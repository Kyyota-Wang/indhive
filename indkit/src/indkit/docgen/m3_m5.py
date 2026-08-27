"""Modules 3, 4 and 5, and the closing appendices."""
from __future__ import annotations

from docx.shared import Pt

from .common import (DRUG, GAPS, INVARIANTS, PROTOCOL, REFERENCES, SPONSOR, TRACE,
                    assume, cited, filed, gap, h, note, pagebreak, table,
                    table_source, trace)


def _rewrite_note(doc, module: str, sources: str):
    p = doc.add_paragraph()
    r = p.add_run(f"{module} has been authored from the underlying data rather than transcribed.")
    r.italic = True
    doc.add_comment(
        runs=[r],
        text=(f"Note on scope: 13_IND_Drafting_Instructions §3 directs that {sources} be filed as-is and "
              f"not rewritten. The sponsor has instead instructed that this module be authored afresh. "
              f"This is a deliberate departure from the drafting instructions. The source dossier remains "
              f"the controlling document; where this text and the source differ, the source governs."),
        author="IRIS", initials="IR")


# ================================================================= MODULE 3

def build_m3(doc):
    h(doc, 1, "Module 3 — Quality")
    _rewrite_note(doc, "Module 3", "07_CMC_Dossier_PMX103_v3.docx")

    note(doc,
         "A note on the source: the CMC dossier contains 478 substantive paragraphs of which 94 are "
         "distinct. Two passages — one describing the graded Phase 1 approach and one describing the "
         "analytical methods — are repeated 26 and 24 times respectively across different subsections. "
         "The summary below is built from the distinct content. Sections whose only content in the "
         "source is repeated boilerplate are marked as not established rather than paraphrased into "
         "something that appears substantive. This is raised for the sponsor's attention.")

    # -------------------------------------------------------------- 3.2.S
    h(doc, 2, "3.2.S Drug Substance")

    h(doc, 3, "3.2.S.1 General Information")
    cited(doc,
          "PMX-103 is a recombinant humanised bispecific T-cell engager of approximately 55 kDa that "
          "binds the PRAME-derived peptide ALYVDSLFFL, corresponding to PRAME residues 425–433 and "
          "presented on HLA-A*02:01, and CD3ε on human T cells. No INN has been assigned.",
          "07_CMC 3.2.S.1.1–3.2.S.1.2 — GAP-03")
    cited(doc,
          "The molecule is highly soluble in aqueous buffer across pH 5.0 to 8.0, adopts a predominantly "
          "β-sheet secondary structure consistent with the immunoglobulin variable domain fold, and has "
          "a melting temperature of approximately 68 °C. These properties give latitude in buffer "
          "selection and support a liquid presentation.",
          "07_CMC 3.2.S.1.3")

    h(doc, 3, "3.2.S.2 Manufacture")
    cited(doc,
          "The drug substance is expressed in a CHO-K1-derived cell line stably transfected to produce "
          "the single-chain bispecific polypeptide, cultured in chemically defined, animal-component-free "
          "medium.",
          "07_CMC 3.2.S.2.2")
    cited(doc,
          "The process comprises thaw of the working cell bank and expansion in shake flasks; seed train "
          "expansion in stirred-tank bioreactors; production culture in a 25 L single-use stirred-tank "
          "bioreactor operated in fed-batch mode for 14 days; and harvest by depth filtration followed "
          "by sterile filtration.",
          "07_CMC 3.2.S.2.2")
    cited(doc,
          "Purification uses orthogonal unit operations: Protein L affinity capture exploiting the "
          "variable kappa light chain, low-pH viral inactivation, cation-exchange and anion-exchange "
          "chromatography to remove process- and product-related impurities, 35 kDa nanofiltration, and "
          "a final ultrafiltration and diafiltration step for concentration and buffer exchange.",
          "07_CMC 3.2.S.2.2")
    cited(doc,
          "In-process control monitors viable cell density, viability, glucose, lactate, pH, dissolved "
          "oxygen and product titre. The clarified harvest is tested for bioburden, mycoplasma by PCR "
          "and adventitious virus by in vitro assay before purification proceeds.",
          "07_CMC 3.2.S.2.4")
    cited(doc,
          "Process targets are at least 95 percent monomer by size-exclusion chromatography, host cell "
          "protein no greater than 100 ng/mg, and host cell DNA below the stated limit.",
          "07_CMC 3.2.S.2.2")

    gap(doc,
        "The manufacturing site is not identified.",
        "3.2.S.2.1 of the source dossier contains only the general statement about the graded Phase 1 "
        "approach; it names no manufacturer, address or FEI number. A Phase 1 IND is expected to "
        "identify the facility at which the drug substance is made.",
        "Which facility manufactures the drug substance, and at what address and FEI number? Is "
        "manufacture performed in-house or at a contract manufacturer?",
        "CMC (Dr. Kim)")

    h(doc, 3, "3.2.S.3 Characterisation")
    cited(doc,
          "Size-exclusion chromatography quantifies monomer content together with high- and "
          "low-molecular-weight species and is the primary control for aggregation and fragmentation. "
          "Imaged capillary isoelectric focusing quantifies charge variants as an indicator of process "
          "consistency and chemical modification. A cell-based assay measures target-dependent lysis as "
          "the potency attribute.",
          "07_CMC 3.2.S.3.1–3.2.S.3.2")

    h(doc, 3, "3.2.S.4 Control of Drug Substance")
    cited(doc,
          "Analytical procedures are validated in accordance with ICH Q2(R1). The clinical drug "
          "substance lot is PMX-103-DS-25-017; the lot used for the toxicology programme is "
          "PMX-103-DS-25-014.",
          "07_CMC 3.2.S.4.2–3.2.S.4.4; 14 Invariants")
    gap(doc,
        "Batch analysis data are illustrative.",
        "The input package states that batch values in the CMC dossier are illustrative and that actual "
        "certificates of analysis have not been issued.",
        "Are release certificates of analysis available for PMX-103-DS-25-017 and PMX-103-DS-25-014?",
        "CMC (Dr. Kim)", "GAP-08")

    h(doc, 3, "3.2.S.7 Stability")
    filed(doc, "Drug substance stability summary, protocol and data", "07_CMC 3.2.S.7",
          "No shelf-life or storage conclusion is restated here; the input brief does not extract one.")

    # -------------------------------------------------------------- 3.2.P
    h(doc, 2, "3.2.P Drug Product")

    h(doc, 3, "3.2.P.1 Description and Composition")
    cited(doc,
          "The drug product is a sterile solution for intravenous infusion, 1.0 mg in 1.1 mL, in a 2 mL "
          "Type I glass vial closed with a 13 mm butyl stopper and a flip-off seal. It is stored at "
          "2–8 °C protected from light and is single-use.",
          "07_CMC 3.2.P.1; 12 Section D.6")

    h(doc, 3, "3.2.P.5 Control of Drug Product")
    doc.add_paragraph("The release specification is given in full at Module 2.3.P.5. The clinical drug "
                      "product lot is PMX-103-DP-26-007.")
    table_source(doc, "07_CMC 3.2.P.5.1, 3.2.P.5.4; 12 Section D.6")

    h(doc, 3, "3.2.P.7 Stability")
    filed(doc, "Drug product stability summary, protocol and data", "07_CMC 3.2.P.7")

    gap(doc,
        "Excipient composition is not stated.",
        "The release specification gives strength, container and storage but not the formulation buffer, "
        "tonicity agent, surfactant or stabiliser. 3.2.P.1 and 3.2.P.4 of the source dossier do not "
        "state a quantitative composition. A quantitative formulation is required for the IND.",
        "What is the quantitative composition of the drug product per vial — buffer species and "
        "molarity, pH, tonicity agent, surfactant and any stabiliser?",
        "CMC (Dr. Kim)")

    # -------------------------------------------------------------- 3.2.A / R
    h(doc, 2, "3.2.A Appendices")
    h(doc, 3, "3.2.A.2 Adventitious Agents Safety Evaluation")
    cited(doc,
          "Viral clearance was evaluated in a scaled-down study in accordance with ICH Q5A(R2), using "
          "the enveloped viruses xenotropic murine leukaemia virus and pseudorabies virus and the "
          "non-enveloped viruses minute virus of mice and reovirus type 3, across two orthogonal steps: "
          "low-pH inactivation and 35 kDa nanofiltration. Total log reduction for enveloped virus was at "
          "least 8.0.",
          "07_CMC 3.2.A.2")
    assume(doc,
           "The use of chemically defined, animal-component-free medium is taken to mean that no "
           "animal-derived raw material enters the process.",
           "07_CMC 3.2.S.2.2, 3.2.A.2",
           assumption=(
               "Assumption: no animal-derived materials are used anywhere in the process, including in cell "
               "bank establishment. The dossier describes the production medium as animal-component-free but "
               "does not make the same statement about the original cell bank. If bovine or other animal "
               "material was used historically in banking, a TSE/BSE statement is required."),
           anchor="no animal-derived raw material enters the process")

    h(doc, 2, "3.2.R Regional Information")
    filed(doc, "Executed batch records, comparability protocols and methods validation package",
          "07_CMC 3.2.R.1–3.2.R.3")

    trace("M3", "Quality", "07_CMC v3.0 (deduplicated)", "Rewritten, 1 assumption, 3 gaps")
    pagebreak(doc)


# ================================================================= MODULE 4

def build_m4(doc):
    h(doc, 1, "Module 4 — Nonclinical Study Reports")
    _rewrite_note(doc, "Module 4",
                  "09_Nonclinical_Development_Plan_PMX103_v3.docx and 11_Toxicology_Dossier_PMX103_v3.docx")

    h(doc, 2, "4.2.1 Pharmacology Study Reports")
    h(doc, 3, "4.2.1.1 Primary Pharmacodynamics")
    cited(doc,
          "Redirected lysis of PRAME-positive, HLA-A*02:01-positive targets was characterised in vitro "
          "with human peripheral blood mononuclear cells, and anti-tumour activity was demonstrated in "
          "PRAME-positive xenograft models in the HHD-DR1 transgenic mouse, with complete regression in "
          "all animals at 10 μg/kg. Parameters are tabulated at Module 2.6.3.",
          "11_TOX 4.2.1.1; 12 Sections D.2, D.4")
    h(doc, 3, "4.2.1.2 Secondary Pharmacodynamics")
    cited(doc, "Secondary pharmacology screening identified no off-target binding of concern.",
          "11_TOX 4.2.1.2")
    h(doc, 3, "4.2.1.3 Safety Pharmacology")
    cited(doc,
          "Three studies were conducted: cardiovascular telemetry in the cynomolgus monkey "
          "(PMX-103-SP-001), respiratory plethysmography in the rat (PMX-103-SP-002) and an Irwin "
          "central nervous system battery in the rat (PMX-103-SP-003). Results are tabulated at Module "
          "2.6.7.",
          "11_TOX 4.2.1.3; 12 Section D.5")

    h(doc, 2, "4.2.2 Pharmacokinetics Study Reports")
    cited(doc,
          "Reports cover the bioanalytical method, pharmacokinetics in mice and in cynomolgus monkeys, "
          "toxicokinetics accompanying the repeat-dose studies, and the human pharmacokinetic "
          "projection by allometric scaling. Parameters are tabulated at Module 2.6.5.",
          "11_TOX 4.2.2.1–4.2.2.5; 12 Section D.3")

    h(doc, 2, "4.2.3 Toxicology Study Reports")
    h(doc, 3, "4.2.3.1 Species Selection")
    cited(doc,
          "The HLA-A2.1 transgenic HHD-DR1 mouse and the cynomolgus monkey were selected as "
          "pharmacologically relevant species. The wild-type mouse lacks the human restriction element "
          "and was used only for passive disposition.",
          "11_TOX 4.2.3.1")
    h(doc, 3, "4.2.3.2 – 4.2.3.6 Single- and Repeat-Dose Toxicity")
    cited(doc,
          "Six studies were conducted, comprising single-dose studies in wild-type and transgenic mice "
          "and in the cynomolgus monkey, repeat-dose studies with 28-day recovery in both species, and "
          "one GLP repeat-dose study in the transgenic mouse. Designs, doses and key findings are "
          "tabulated at Module 2.6.7 and the derived safety values at Module 2.4.4.",
          "11_TOX 4.2.3.2–4.2.3.6; 12 Section D.4")
    h(doc, 3, "4.2.3.7 Genetic Toxicology and Carcinogenicity")
    cited(doc,
          "Not conducted, consistent with ICH S6(R1) and ICH S9 for a protein therapeutic in advanced "
          "malignancy. See the assumption recorded at Module 2.4.4.",
          "11_TOX 4.2.3.7")
    h(doc, 3, "4.2.3.8 Reproductive and Developmental Toxicology")
    cited(doc, "Not conducted, on the same basis.", "11_TOX 4.2.3.8")
    h(doc, 3, "4.2.3.9 Immunogenicity")
    cited(doc,
          "Anti-drug antibodies were detected in two of eight cynomolgus monkeys at day 28 of the "
          "repeat-dose study PMX-103-TOX-005.",
          "11_TOX 4.2.3.9; 12 Section D.4")
    trace("M4", "Nonclinical study reports", "11_TOX v3.0; 09_NCD v3.0", "Rewritten")
    pagebreak(doc)


# ================================================================= MODULE 5

def build_m5(doc):
    h(doc, 1, "Module 5 — Clinical Study Reports")
    h(doc, 2, "5.2 Tabular Listing of All Clinical Studies")
    cited(doc,
          f"No clinical study of {DRUG} has been conducted or is ongoing. Protocol {PROTOCOL} is the "
          "first study proposed and has not started.",
          "12 Section B.9 — GAP-07")
    h(doc, 2, "5.3 Clinical Study Reports")
    gap(doc,
        "No clinical study reports exist.",
        "The Phase 1 study has not begun, so there are no reports of biopharmaceutic studies, human "
        "pharmacokinetic studies, efficacy or safety studies to file.",
        "None — expected for an original IND. See GAP-07.",
        "Clinical", "GAP-07")
    h(doc, 2, "5.3.5.2 Study Reports of Uncontrolled Clinical Studies — Protocol")
    filed(doc, f"The clinical protocol {PROTOCOL} (ICH E6(R2) Section 6 layout, sections 1–16 with the "
               "Schedule of Assessments)",
          "10_Phase1_Protocol_PMX103_P1_001_v3.docx",
          "Objectives, eligibility, the mTPI-2 design and decision table, DLT definitions, ASTCT-based "
          "CRS and ICANS management, stopping rules and the statistical analysis plan.")
    gap(doc,
        "The protocol version to be filed is not final.",
        "The package supplies version 3.0 dated 22 August 2026 and records that finalisation is "
        "outstanding. The version filed with the IND must be the executed version.",
        "Is v3.0 dated 22 August 2026 the version to be filed, or is a later version expected before "
        "submission?",
        "Clinical", "GAP-09")
    trace("M5", "Clinical study reports", "10_Protocol v3.0; GAP-07", "None — study not started")
    pagebreak(doc)


# ================================================================= APPENDICES

def build_appendices(doc):
    h(doc, 1, "Appendix A — Traceability")
    doc.add_paragraph(
        "Every section of this document, the source it was written from, and how it was completed. "
        "A reviewer wanting to check the document rather than read it should start here.")
    table(doc, ["Section", "Title", "Source", "Status"],
          [[s, t, src, st] for s, t, src, st in TRACE],
          widths=[0.9, 2.0, 2.2, 1.5])

    h(doc, 1, "Appendix B — Open Items and Owners")
    doc.add_paragraph(
        "Items carried from the sponsor's gaps log, plus items this drafting exercise identified. "
        "No value has been substituted for any of them.")
    rows = [[g, item, owner, need, "Sponsor log"] for g, item, owner, need in GAPS]
    rows += [
        ["NEW-01", "Reviewing division and CDER address for the cover letter", "Regulatory",
         "Division assignment", "Identified here"],
        ["NEW-02", "Pre-IND meeting status and any FDA commitments", "Regulatory",
         "Meeting minutes or confirmation none held", "Identified here"],
        ["NEW-03", "Investigational drug labelling (21 CFR 312.6)", "CMC",
         "Draft container and carton labels", "Identified here"],
        ["NEW-04", "Drug substance manufacturing site, address and FEI", "CMC",
         "Facility identification", "Identified here"],
        ["NEW-05", "Quantitative drug product composition (excipients)", "CMC",
         "Formulation record", "Identified here"],
        ["NEW-06", "Exposure-based safety margin at the cynomolgus NOAEL", "DMPK / Safety",
         "AUC at 5.0 μg/kg", "Identified here"],
    ]
    table(doc, ["ID", "Item", "Owner", "What would resolve it", "Origin"], rows,
          widths=[0.7, 2.3, 1.2, 1.7, 0.9])

    h(doc, 1, "Appendix C — Invariant Values Used")
    doc.add_paragraph(
        "The sponsor's traceability matrix declares parameters that must read identically wherever they "
        "appear. Every value below is reproduced exactly as declared; none has been rounded, converted "
        "or restated.")
    table(doc, ["Parameter", "Value", "Unit", "Declared source"],
          [[k.replace("_", " "), v, u, s] for k, (v, u, s) in INVARIANTS.items()],
          widths=[1.9, 0.9, 1.1, 2.5])
    table_source(doc, "14_IND_Section_Traceability_Matrix.xlsx — Invariants sheet")

    h(doc, 1, "Appendix D — Observations on the Source Package")
    doc.add_paragraph(
        "Three matters were noticed while drafting. None affects the science; all three affect the "
        "filing, and all are offered for the sponsor's judgement rather than asserted as errors.")
    table(doc,
          ["Observation", "Where", "Why it matters"],
          [["The completed Form 1571 draft numbers its fields 1–17, which does not match the box "
            "numbering on the current form. Item 2 is Box 3, item 6 is Box 7, item 7 is Box 8, item 9 "
            "is Box 11, and item 10 (lettered a–k) is Box 12, which is numbered 1–10.",
            "06_Form_FDA_1571_AutoFilled.docx",
            "Transcription into the official form would misplace content."],
           ["The traceability matrix maps 312.23(a)(7) to \"M1.7 / M3\" and (a)(8) to "
            "\"M1.8 / M2.6 / M4\". M1.7 and M1.8 are not eCTD locations. Section C of the input brief "
            "gives the correct mapping.",
            "14 Section Map sheet vs 12 Section C",
            "The two indexes disagree with each other; a compiler following the matrix would misfile."],
           ["The nonclinical development plan numbers its summaries M2.6.2 pharmacology written, "
            "M2.6.3 pharmacokinetics written, M2.6.4 toxicology written. Under ICH M4S, 2.6.3 is the "
            "pharmacology tabulated summary and pharmacokinetics written is 2.6.4.",
            "09_NCD section headings",
            "Non-standard numbering in Module 2 is a common validation finding."],
           ["The CMC dossier contains 478 substantive paragraphs of which 94 are distinct; two "
            "passages repeat 26 and 24 times across different subsections. Several sections, including "
            "3.2.S.2.1 and 3.2.P.3.1, contain only the repeated text.",
            "07_CMC_Dossier_PMX103_v3.docx",
            "Sections that appear complete are empty of section-specific content."]],
          widths=[3.0, 1.5, 2.0])

    h(doc, 1, "Appendix E — References")
    doc.add_paragraph(
        "Cited as listed in the sponsor's drafting instructions. No title or date has been altered and "
        "no reference has been added.")
    for title, url in REFERENCES:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(title)
        r = p.add_run(f" — {url}")
        r.font.size = Pt(8)
    table_source(doc, "13_IND_Drafting_Instructions_PMX103.docx §4 and References")

"""Shared building blocks for the PMX-103 IND draft.

Four conventions carry the whole document, and every one of them is visible to a
reader without a legend:

  cited(...)   plain text that ends in a bracketed source. Nothing is asserted
               without naming the document it came from.
  assume(...)  plain text plus a real Word comment authored by IRIS, stating the
               assumption the sentence rests on. The reader can accept, reject or
               reply to each one in Word.
  gap(...)     blue text. The fact is not available; the paragraph says why, names
               the owner, and asks the question that would unblock it.
  filed(...)   a cross-reference, not a gap. The content exists in a source dossier
               that is filed with the application.

The distinction between gap() and filed() matters: without it a correct document
looks like it has a hundred failures.
"""
from __future__ import annotations

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches

BLUE = RGBColor(0x00, 0x00, 0xCC)
GREY = RGBColor(0x59, 0x59, 0x59)
AUTHOR = "IRIS"
INITIALS = "IR"


# ---------------------------------------------------------------- text blocks

def cited(doc, text: str, source: str, style: str | None = None):
    """A factual statement with its source. The default for ordinary content."""
    p = doc.add_paragraph(style=style)
    p.add_run(text.rstrip() + " ")
    s = p.add_run(f"[{source}]")
    s.font.size = Pt(9)
    s.font.color.rgb = GREY
    return p


def assume(doc, text: str, source: str, assumption: str, anchor: str | None = None):
    """A statement that is only sound if an assumption holds.

    The assumption becomes a Word comment anchored to the specific clause it
    applies to, so the reviewer sees what is being taken for granted and where.
    """
    p = doc.add_paragraph()
    if anchor and anchor in text:
        head, _, tail = text.partition(anchor)
        if head:
            p.add_run(head)
        target = p.add_run(anchor)
        if tail:
            p.add_run(tail.rstrip() + " ")
    else:
        target = p.add_run(text.rstrip() + " ")
    s = p.add_run(f"[{source}]")
    s.font.size = Pt(9)
    s.font.color.rgb = GREY
    doc.add_comment(runs=[target], text=assumption, author=AUTHOR, initials=INITIALS)
    return p


def gap(doc, heading: str, reason: str, question: str, owner: str, gap_id: str | None = None):
    """Content that cannot be written, in blue, with the question that unblocks it."""
    p = doc.add_paragraph()
    tag = f"[NOT COMPLETED{' · ' + gap_id if gap_id else ''}] "
    r = p.add_run(tag + heading)
    r.bold = True
    r.font.color.rgb = BLUE

    body = doc.add_paragraph()
    b = body.add_run(f"Why this is not filled in: {reason}")
    b.font.color.rgb = BLUE

    q = doc.add_paragraph()
    qr = q.add_run(f"Question for the sponsor: {question}")
    qr.font.color.rgb = BLUE
    qr.italic = True

    o = doc.add_paragraph()
    orr = o.add_run(f"Owner: {owner}. This placeholder must be resolved before submission; "
                    f"no value has been substituted.")
    orr.font.color.rgb = BLUE
    orr.font.size = Pt(9)
    return p


def filed(doc, what: str, where: str, note: str | None = None):
    """Content that lives in a source dossier filed with the application."""
    p = doc.add_paragraph()
    r = p.add_run(f"{what} — filed in full at {where}.")
    r.italic = True
    if note:
        n = p.add_run(f" {note}")
        n.italic = True
        n.font.size = Pt(9)
        n.font.color.rgb = GREY
    return p


def note(doc, text: str):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(9)
    r.font.color.rgb = GREY
    r.italic = True
    return p


# ---------------------------------------------------------------- structure

def h(doc, level: int, text: str):
    return doc.add_heading(text, level=level)


def table(doc, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, head in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(head)
        run.bold = True
        run.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
    if widths:
        for r in t.rows:
            for i, w in enumerate(widths):
                r.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t


def table_source(doc, source: str):
    p = doc.add_paragraph()
    r = p.add_run(f"Source: {source}")
    r.font.size = Pt(8.5)
    r.font.color.rgb = GREY
    return p


def pagebreak(doc):
    doc.add_page_break()


# ---------------------------------------------------------------- fixed data

SPONSOR = "Indela Therapeutics, Inc."
SPONSOR_ADDR = "1515 Hollins Street, Suite 400, Baltimore, MD 21202"
DRUG = "PMX-103"
DRUG_FULL = ("PMX-103 — recombinant humanised anti-PRAME/HLA-A*02:01-peptide × "
             "anti-CD3ε bispecific T-cell engager")
PROTOCOL = "PMX-103-P1-001"
SUBMIT_DATE = "21 August 2026"

# 14_IND_Section_Traceability_Matrix.xlsx, Invariants sheet. These values are
# authoritative; nothing in this document may restate them differently.
INVARIANTS = {
    "starting_dose": ("0.5", "μg flat", "Master data sheet"),
    "starting_cmax": ("0.041", "ng/mL", "PK projection"),
    "ec10": ("0.12", "ng/mL", "PMX-103-IVT-002"),
    "ec50": ("0.62", "ng/mL", "PMX-103-IVT-002"),
    "mabel": ("1.0", "μg/kg", "PMX-103-TOX-002"),
    "noael_cyno": ("5.0", "μg/kg", "PMX-103-TOX-005"),
    "noael_mouse": ("10.0", "μg/kg", "PMX-103-TOX-006"),
    "mtd_mouse": ("30.0", "μg/kg", "PMX-103-TOX-003"),
    "hnstd_cyno": ("50.0", "μg/kg", "PMX-103-TOX-005"),
    "cyno_cmax": ("34.2", "ng/mL", "PK cyno 40 μg/kg"),
    "cyno_auc": ("1281", "ng·h/mL", "PK cyno"),
    "human_cl": ("0.21", "mL/h/kg", "Allometric scaling"),
    "human_vss": ("0.18", "L/kg", "Allometric scaling"),
    "qtcf": ("9.2", "ms", "PMX-103-SP-001"),
    "dlt_window": ("28", "days", "Protocol 8.3"),
}

GAPS = [
    ("GAP-01", "Actual IND number (placeholder 100,001 reserved)", "Regulatory (Whitfield)", "FDA IND receipt"),
    ("GAP-02", "Form 1571 signatory confirmation", "Regulatory (Whitfield)", "Sponsor authorisation"),
    ("GAP-03", "Proprietary / INN name (not assigned)", "Regulatory", "WHO INN / trademark"),
    ("GAP-04", "Site and investigator list (Form 1572)", "Clinical (Tanaka)", "Site selection"),
    ("GAP-05", "ClinicalTrials.gov NCT ID (Form 3674)", "Clinical (Tanaka)", "ClinicalTrials.gov registration"),
    ("GAP-06", "Financial disclosure Forms 3454/3455", "Regulatory", "Per-investigator"),
    ("GAP-07", "Clinical data (none — Phase 1 not started)", "Clinical", "n/a"),
    ("GAP-08", "Actual batch COAs (illustrative values only)", "CMC (Kim)", "QC certificates"),
    ("GAP-09", "Final protocol version date (v3.0 dated 22 August 2026)", "Clinical", "Protocol finalisation"),
    ("GAP-10", "IRB approvals at each site", "Clinical", "Site IRBs"),
]

CONTRIBUTORS = [
    ("Dr. Elena Vasquez, PhD", "Vice President, Discovery Biology", "Discovery Biology"),
    ("Dr. Marcus Chen, PhD", "Director, DMPK & Clinical Pharmacology", "DMPK / Clinical Pharmacology"),
    ("Dr. Priya Raman, DVM, PhD, DABT", "Senior Director, Toxicology & Safety Assessment", "Safety Assessment"),
    ("Dr. James Okafor, PhD", "Director, Nonclinical Pharmacology", "Dose-Range Finding"),
    ("Dr. Sofia Lindgren, PhD", "Director, Translational Sciences", "Translational Sciences"),
    ("Dr. Robert Tanaka, MD, PhD", "VP, Clinical Development & Medical Affairs", "Clinical Development"),
    ("Dr. Anita Desai, PhD", "Director, Biostatistics", "Biostatistics"),
    ("Ms. Laura Whitfield, JD, RAC", "VP, Regulatory Affairs", "Regulatory Affairs"),
    ("Dr. David Kim, PhD", "VP, CMC & Process Sciences", "CMC / Process Sciences"),
    ("Dr. Rachel Osei, MD", "Clinical Pharmacologist", "Clinical Pharmacology"),
]

# 13_IND_Drafting_Instructions §4: cite these as listed; do not re-title or re-date.
REFERENCES = [
    ("21 CFR 312.23 — IND content and format", "https://www.law.cornell.edu/cfr/text/21/312.23"),
    ("FDA — Content and Format of INDs for Phase 1 Studies (Nov 1995)", "https://www.fda.gov/media/71203/download"),
    ("FDA — Bispecific Antibody Development Programs Guidance (May 2024)", "https://www.fda.gov/media/123313/download"),
    ("ICH S9 — Nonclinical Evaluation for Anticancer Pharmaceuticals", "https://www.fda.gov/media/73161/download"),
    ("ICH M3(R2) — Nonclinical Safety Studies for the Conduct of Human Clinical Trials", "https://www.fda.gov/media/71542/download"),
    ("ICH S6(R1) — Preclinical Safety Evaluation of Biotechnology-Derived Pharmaceuticals",
     "https://www.fda.gov/regulatory-information/search-fda-guidance-documents/s6r1-preclinical-safety-evaluation-biotechnology-derived-pharmaceuticals"),
    ("ICH E6(R2) — Good Clinical Practice: Integrated Addendum", "https://www.fda.gov/media/169090/download"),
    ("FDA — Estimating the Maximum Safe Starting Dose in Initial Clinical Trials (July 2005)", "https://www.fda.gov/media/72309/download"),
    ("FDA — QSP-Based Dose Selection / MABEL in FIH (draft)", "https://www.fda.gov/media/193230/download"),
    ("FDA — Project Optimus", "https://www.fda.gov/about-fda/oncology-center-excellence/project-optimus"),
    ("FDA — Phase 1 IND Navigator", "https://www.fda.gov/industry/phase-1-investigational-new-drug-ind-navigator"),
    ("Form FDA 1571", "https://www.fda.gov/media/77596/download"),
    ("Form FDA 1572", "https://www.fda.gov/media/71816/download"),
    ("Form FDA 3674 / IND forms landing", "https://www.fda.gov/drugs/investigational-new-drug-ind-application/ind-forms-and-instructions"),
    ("Chang AY et al. — Pr20 TCRm anti-PRAME/HLA-A2 (Nat Med 2017)", "https://pmc.ncbi.nlm.nih.gov/articles/PMC5490756/"),
    ("Immatics IMA402 TCER (NCT05958121)",
     "https://investors.immatics.com/news-releases/news-release-details/immatics-initiates-phase-12-clinical-trial-evaluate-prame-tcr"),
    ("Wermke M et al. — IMA203 PRAME TCR-T (Nat Med 2025)", "https://www.nature.com/articles/s41591-025-03650-6"),
    ("Skrzypczynska K et al. — PRAME pMHC TCE design (mAbs 2025)", "https://www.tandfonline.com/doi/full/10.1080/19420862.2025.2563773"),
    ("Guo W, Yuan Y, Ji Y — mTPI-2 design (R package escalation, get_mtpi2)",
     "https://search.r-project.org/CRAN/refmans/escalation/html/get_mtpi2.html"),
    ("Lee DW et al. — ASTCT CRS and ICANS consensus (2019)", "https://www.astctjournal.org"),
]

# Traceability rows are appended by each module builder and rendered in Appendix B.
TRACE: list[tuple[str, str, str, str]] = []


def trace(section: str, title: str, source: str, status: str):
    TRACE.append((section, title, source, status))

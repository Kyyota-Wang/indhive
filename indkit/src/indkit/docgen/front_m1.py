"""Front matter and Module 1 — Administrative Information."""
from __future__ import annotations

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from .common import (BLUE, GREY, CONTRIBUTORS, DRUG, DRUG_FULL, GAPS, PROTOCOL,
                    SPONSOR, SPONSOR_ADDR, SUBMIT_DATE, assume, cited, filed,
                    gap, h, note, pagebreak, table, table_source, trace)


def build_front(doc):
    # ------------------------------------------------------------- title page
    t = doc.add_heading("INVESTIGATIONAL NEW DRUG APPLICATION", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for line, size, bold in [
        (DRUG_FULL, 13, True),
        (f"Protocol {PROTOCOL} — Phase 1, first-in-human", 11, False),
        (SPONSOR, 11, False),
        (SPONSOR_ADDR, 10, False),
        (f"Original submission (Serial Number 0000) — {SUBMIT_DATE}", 10, False),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.size = Pt(size)
        r.bold = bold

    doc.add_paragraph()
    box = doc.add_paragraph()
    box.alignment = WD_ALIGN_PARAGRAPH.CENTER
    w = box.add_run(
        "DRAFTING EXERCISE — NOT A REGULATORY SUBMISSION\n"
        "PMX-103 and Indela Therapeutics, Inc. are fictional. This document was assembled by an "
        "AI drafting agent from a synthetic input package to evaluate machine-authored IND content. "
        "It has not been reviewed or approved by any sponsor and must not be filed with FDA or "
        "presented as a genuine application."
    )
    w.bold = True
    w.font.color.rgb = RGBColor(0x99, 0x2D, 0x2D)
    w.font.size = Pt(10)

    pagebreak(doc)

    # --------------------------------------------------------- document control
    h(doc, 1, "Document Control")
    table(doc,
          ["Field", "Value"],
          [["Document", f"Investigational New Drug Application — {DRUG}"],
           ["Compound", DRUG],
           ["Protocol", PROTOCOL],
           ["Sponsor", SPONSOR],
           ["Submission type", "Original IND (Serial Number 0000)"],
           ["Submission date", SUBMIT_DATE],
           ["IND number", "Not assigned — see GAP-01"],
           ["Drafted by", "IRIS (AI drafting agent), from the input package of 22 August 2026"],
           ["Status", "Draft for sponsor review — not for submission"]],
          widths=[1.7, 4.6])

    h(doc, 1, "How to Read This Document")
    doc.add_paragraph(
        "Four conventions run through the whole document. They exist so that a reviewer can tell, "
        "at any point, how much weight a sentence carries and who has to act on it."
    )
    table(doc,
          ["Convention", "Meaning", "What the reviewer should do"],
          [["Plain text with a bracketed source",
            "A fact taken from the input package. The bracket names the document it came from.",
            "Spot-check against the cited source."],
           ["Word comment by IRIS",
            "The statement is sound only if an assumption holds. The comment states the assumption.",
            "Accept, correct or reject each comment in Word."],
           ["Blue text",
            "The content could not be written. The paragraph says why, names the owner and asks the "
            "question that would unblock it. No value has been substituted.",
            "Answer the question, or confirm the gap is expected."],
           ["Italic cross-reference",
            "The content exists in a source dossier filed with this application. It is not missing.",
            "Nothing — it is a pointer, not a gap."]],
          widths=[1.5, 3.0, 1.8])
    note(doc,
         "The distinction between blue text and an italic cross-reference is deliberate. Without it a "
         "document that is behaving correctly would appear to contain a hundred failures.")

    h(doc, 1, "Contributors to the Source Package")
    doc.add_paragraph(
        "The input package carries functional authorship for each area of content. The bylines below "
        "are reproduced from the source documents; this drafting agent authored no scientific content "
        "of its own."
    )
    table(doc, ["Name", "Title", "Function"],
          [[n, t_, f] for n, t_, f in CONTRIBUTORS], widths=[2.1, 2.6, 1.7])
    table_source(doc, "12_IND_Drafting_Input_Brief_PMX103.docx; 13_IND_Drafting_Instructions_PMX103.docx")

    # --------------------------------------------------------------- crosswalk
    h(doc, 1, "Statutory Content Crosswalk")
    doc.add_paragraph(
        "21 CFR 312.23(a) lists the content an IND must contain. The CTD organises the same material "
        "into Modules 1 to 5. The two are different indexes over the same submission, and conflating "
        "them is a common source of misfiled content. This document is organised by CTD module; the "
        "table below shows where each statutory paragraph is satisfied."
    )
    table(doc,
          ["21 CFR 312.23(a)", "Required content", "Location in this document"],
          [["(a)(1)", "Cover sheet — Form FDA 1571", "M1.1.1"],
           ["(a)(2)", "Table of contents", "M1 index and CTD table of contents (M2.1)"],
           ["(a)(3)", "Introductory statement and general investigational plan", "M1.20"],
           ["(a)(4)", "[Reserved]", "No content required"],
           ["(a)(5)", "Investigator's brochure", "M1.14.4.1 (filed); summarised in M2.7"],
           ["(a)(6)", "Protocols", "M5.3.5.2 (filed); synopsis in M2.5"],
           ["(a)(7)", "Chemistry, manufacturing and control information", "M3.2.S / M3.2.P; summary in M2.3"],
           ["(a)(8)", "Pharmacology and toxicology information", "M4.2; overview in M2.4; summaries in M2.6"],
           ["(a)(9)", "Previous human experience", "M2.5 and M1.20; none for PMX-103"],
           ["(a)(10)", "Additional information", "M1.3, M1.14"],
           ["(a)(11)", "Relevant information", "Appendix B, traceability"]],
          widths=[1.0, 3.1, 2.2])
    note(doc,
         "The input package's traceability matrix maps (a)(7) to \"M1.7 / M3\" and (a)(8) to "
         "\"M1.8 / M2.6 / M4\". M1.7 and M1.8 are not eCTD locations; Module 1 has no such sections. "
         "The mapping used here places quality content in Module 3 and nonclinical content in Module 4, "
         "which is what Section C of the input brief itself specifies. This discrepancy is raised for "
         "the sponsor's attention.")
    trace("Front", "Statutory crosswalk", "21 CFR 312.23(a); 12 Section C", "Authored")

    pagebreak(doc)


# ================================================================= MODULE 1

def build_m1(doc):
    h(doc, 1, "Module 1 — Administrative Information")

    # ------------------------------------------------------------- 1.1 forms
    h(doc, 2, "1.1 Forms")

    h(doc, 3, "1.1.1 Form FDA 1571")
    doc.add_paragraph(
        "The fields below are presented against the box numbers printed on the current Form FDA 1571. "
        "Values are transcribed from the sponsor's completed draft and from the master data sheet."
    )
    table(doc,
          ["Box", "Field", "Value"],
          [["1", "Name of sponsor", SPONSOR],
           ["2", "Date of submission", SUBMIT_DATE],
           ["3", "Address", SPONSOR_ADDR],
           ["4", "Telephone number", "+1 410 555 0143"],
           ["5", "Name(s) of drug", "PMX-103 (code name). No INN or proprietary name assigned — GAP-03."],
           ["6", "IND number", "Not assigned — GAP-01. Left blank; FDA assigns on receipt."],
           ["7", "Indication(s)",
            "Locally advanced or metastatic solid tumours that are PRAME-positive (central IHC "
            "H-score ≥ 80) and HLA-A*02:01-positive, after ≥ 1 prior line of standard therapy "
            "including a PD-(L)1 inhibitor where indicated."],
           ["8", "Phase(s) of clinical investigation", "Phase 1"],
           ["9", "Numbers of referenced INDs, NDAs, DMFs, BLAs", "None"],
           ["10", "Serial number", "0000"],
           ["11", "This submission contains", "Initial Investigational New Drug Application (IND)"],
           ["12", "Contents of application", "See the checklist below"],
           ["14", "Person responsible for monitoring the conduct and progress of the investigations",
            "Dr. Robert Tanaka, MD, PhD — VP, Clinical Development & Medical Affairs"],
           ["15", "Person(s) responsible for review and evaluation of safety information",
            "Dr. Priya Raman, DVM, PhD, DABT — Senior Director, Toxicology & Safety Assessment; "
            "Dr. Robert Tanaka, MD, PhD"]],
          widths=[0.5, 2.0, 3.8])
    table_source(doc, "06_Form_FDA_1571_AutoFilled.docx; 12 Sections A.1–A.2; Form FDA 1571 (fda.gov/media/77596/download)")

    assume(doc,
           "Box 14 and Box 15 have been assigned to the clinical and safety functions named in the "
           "contributor roster, because the sponsor's draft form does not fill these boxes.",
           "12 Section A.1; contributor roster",
           assumption=(
               "Assumption: Box 14 (monitoring) is assigned to Dr. Tanaka and Box 15 (safety review) to "
               "Dr. Raman with Dr. Tanaka, inferred from their stated functions. The sponsor's draft form "
               "does not name these individuals. If the sponsor has designated different people — a "
               "contract medical monitor, for example — these boxes must be corrected before signature."),
           anchor="assigned to the clinical and safety functions named in the contributor roster")

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("Box 12 — Contents of application")
    r.bold = True
    table(doc,
          ["Item", "Content", "Included"],
          [["1", "Form FDA 1571", "Yes"],
           ["2", "Table of contents", "Yes"],
           ["3", "Introductory statement and general investigational plan", "Yes — M1.20"],
           ["4", "[Reserved]", "—"],
           ["5", "Investigator's brochure", "Yes — filed at M1.14.4.1"],
           ["6", "Protocol(s)", "Yes — filed at M5.3.5.2"],
           ["7", "Chemistry, manufacturing and control data", "Yes — Module 3"],
           ["8", "Pharmacology and toxicology data", "Yes — Module 4"],
           ["9", "Previous human experience", "None for PMX-103; class precedent summarised"],
           ["10", "Additional information", "Yes — M1.3, M1.14"]],
          widths=[0.6, 3.9, 1.8])
    trace("M1.1.1", "Form FDA 1571", "06_Form_FDA_1571_AutoFilled.docx; 12 A.1–A.2", "Authored, 2 assumptions")

    note(doc,
         "The sponsor's completed draft numbers its fields 1 to 17, which does not correspond to the box "
         "numbering on the current Form FDA 1571. Its item 2 (address) is Box 3 on the form, item 6 "
         "(indication) is Box 7, item 7 (phase) is Box 8, item 9 (type of submission) is Box 11, and item "
         "10 (contents, lettered a–k) is Box 12, which is numbered 1 to 10. The transcription above uses "
         "the official box numbers. This is raised for the sponsor's attention.")

    h(doc, 3, "1.1.2 Form FDA 1572 — Statement of Investigator")
    gap(doc,
        "Form FDA 1572 cannot be completed.",
        "A separate Form 1572 is required for each participating investigator, and must carry that "
        "investigator's name, site address, curriculum vitae, clinical laboratory facilities and IRB "
        "information. The input package states that site selection has not been completed, so none of "
        "this information exists yet.",
        "Have any investigators or sites been selected? If site selection is complete, the investigator "
        "names, addresses, IRBs and sub-investigator lists are needed. If not, when is selection expected "
        "to close?",
        "Clinical (Dr. Tanaka)", "GAP-04")
    trace("M1.1.2", "Form FDA 1572", "GAP-04", "Not completed")

    h(doc, 3, "1.1.3 Form FDA 3674 — Certification of Compliance, ClinicalTrials.gov")
    gap(doc,
        "Form FDA 3674 cannot be certified.",
        "The certification requires the ClinicalTrials.gov NCT identifier for the trial, or an explicit "
        "statement that the trial is not an applicable clinical trial. The input package records that "
        "registration has not yet been made.",
        "Has the trial been registered on ClinicalTrials.gov, and if so what is the NCT number? If "
        "registration is deliberately deferred, does the sponsor intend to certify under 42 U.S.C. "
        "282(j)(5)(B) that the trial is not yet an applicable clinical trial?",
        "Clinical (Dr. Tanaka)", "GAP-05")
    trace("M1.1.3", "Form FDA 3674", "GAP-05", "Not completed")

    pagebreak(doc)

    # ------------------------------------------------------- 1.2 cover letter
    h(doc, 2, "1.2 Cover Letter")
    _cover_letter(doc)
    trace("M1.2", "Cover letter", "12 Sections A–B; 10_Protocol", "Authored")

    pagebreak(doc)

    # --------------------------------------------------- 1.3 administrative
    h(doc, 2, "1.3 Administrative Information")

    h(doc, 3, "1.3.1 Contact and Agent Information")
    cited(doc,
          f"All correspondence relating to this application should be directed to Ms. Laura Whitfield, "
          f"JD, RAC, Vice President, Regulatory Affairs, {SPONSOR}, {SPONSOR_ADDR}, telephone "
          f"+1 410 555 0143.",
          "12 Section A.1; contributor roster")
    assume(doc,
           "The sponsor is a domestic corporation and no United States agent is required.",
           "12 Section A.1",
           assumption=(
               "Assumption: no US agent is designated. The sponsor's address is in Baltimore, Maryland, "
               "so a US agent under 21 CFR 312.23(a)(1) is not required. If the sponsor is in fact a "
               "subsidiary of a foreign parent that will hold the IND, an agent must be named."),
           anchor="no United States agent is required")
    trace("M1.3.1", "Contact / agent information", "12 A.1", "Authored, 1 assumption")

    h(doc, 3, "1.3.4 Financial Certification and Disclosure")
    gap(doc,
        "Forms FDA 3454 and 3455 cannot be completed.",
        "Financial certification and disclosure are made per covered clinical investigator. Because no "
        "investigators have been selected (GAP-04), there is no one to certify for and no disclosable "
        "arrangements to report.",
        "Once investigators are selected, will the sponsor certify on Form 3454 that no covered "
        "investigator has a disclosable financial interest, or does it anticipate disclosures on Form "
        "3455 — for example equity held by an academic founder?",
        "Regulatory", "GAP-06")
    trace("M1.3.4", "Financial certification and disclosure", "GAP-06", "Not completed")

    # ------------------------------------------------------- 1.4 references
    h(doc, 2, "1.4 References")
    assume(doc,
           "No letter of authorisation or statement of right of reference is included, because this "
           "application does not rely on a Drug Master File or on another sponsor's application.",
           "12 Section B.11; Form 1571 Box 9",
           assumption=(
               "Assumption: nothing is incorporated by reference. Box 9 of the sponsor's own form lists no "
               "referenced applications, and the CMC dossier describes manufacture at the sponsor's own "
               "process. If any raw material, cell bank or fill-finish operation is covered by a third-party "
               "DMF, a letter of authorisation is required and this section must be completed."),
           anchor="does not rely on a Drug Master File or on another sponsor's application")
    trace("M1.4", "References", "Assumption — no DMF reliance", "Authored, 1 assumption")

    # ------------------------------------------------ 1.12 correspondence
    h(doc, 2, "1.12 Other Correspondence")
    gap(doc,
        "Pre-IND correspondence status is unknown.",
        "The input package does not record whether a pre-IND meeting was requested or held. For a "
        "first-in-human bispecific T-cell engager this is material: FDA advice on the starting dose, the "
        "step-up schedule and the cytokine release syndrome mitigation plan would normally be filed here "
        "and reflected in the protocol.",
        "Was a pre-IND meeting held or a written response received? If so, the meeting request, briefing "
        "package, FDA preliminary responses and minutes should be filed at M1.12.1, and any commitment "
        "made in that exchange should be reflected in the protocol.",
        "Regulatory (Ms. Whitfield)")
    trace("M1.12", "Other correspondence", "Not stated in package", "Not completed")

    # ---------------------------------------------------------- 1.14 labeling
    h(doc, 2, "1.14 Labeling")
    h(doc, 3, "1.14.4.1 Investigator's Brochure")
    filed(doc, "The Investigator's Brochure (IB-PMX-103-002, ICH E6(R2) Section 7 layout)",
          "08_Investigator_Brochure_PMX103_v3.docx",
          "Sections 1–7, including PRAME biology, the ALYVDSLFFL/HLA-A*02:01 epitope, nonclinical "
          "pharmacology and toxicology tables, the human PK projection, and CRS/ICANS grading and "
          "management guidance.")
    trace("M1.14.4.1", "Investigator's brochure", "08_IB v3.0", "Filed as-is")

    h(doc, 3, "1.14.4.2 Investigational Drug Labeling")
    gap(doc,
        "Investigational product container and carton labelling is not available.",
        "21 CFR 312.6 requires investigational labelling bearing the caution statement, and the input "
        "package contains no draft label. The drug product presentation is known — a single-use 2 mL "
        "Type I glass vial containing 1.0 mg in 1.1 mL, stored at 2–8 °C and protected from light — but "
        "the label artwork, lot and expiry format, and caution statement wording are not supplied.",
        "Can draft container and carton labels be provided? If labelling is still in preparation, the "
        "sponsor should confirm the intended caution statement and whether the product will be labelled "
        "for single-patient or multi-patient use.",
        "CMC (Dr. Kim)")
    trace("M1.14.4.2", "Investigational drug labeling", "Not in package", "Not completed")

    pagebreak(doc)

    # ------------------------------------------ 1.20 general investigational plan
    h(doc, 2, "1.20 Introductory Statement and General Investigational Plan")
    _general_plan(doc)
    trace("M1.20", "Introductory statement and general investigational plan",
          "12 B.3; 10_Protocol §§2–4; 09_NCD; 11_TOX", "Authored")

    pagebreak(doc)


# --------------------------------------------------------------- sub-builders

def _cover_letter(doc):
    for line in [SUBMIT_DATE, "", "Food and Drug Administration",
                 "Center for Drug Evaluation and Research", ""]:
        doc.add_paragraph(line)

    gap(doc,
        "The reviewing division and mailing address are not identified.",
        "A cover letter is addressed to the division that will review the application. The input "
        "package does not name one, and the division cannot be inferred reliably: an anti-PRAME "
        "bispecific T-cell engager for solid tumours could be assigned to the Division of Oncology 1 "
        "or to another oncology division depending on the tumour types ultimately enrolled.",
        "Which review division should this be addressed to, and does the sponsor have a CDER document "
        "control room address or an established eCTD gateway route?",
        "Regulatory (Ms. Whitfield)")

    p = doc.add_paragraph()
    p.add_run("Re: ").bold = True
    p.add_run("Original Investigational New Drug Application — Serial Number 0000 — "
              f"{DRUG} — Protocol {PROTOCOL}")

    doc.add_paragraph()
    doc.add_paragraph("Dear Sir or Madam:")
    doc.add_paragraph()

    cited(doc,
          f"{SPONSOR} submits this Original Investigational New Drug Application for {DRUG}, a "
          "recombinant humanised bispecific T-cell engager that binds the PRAME-derived peptide "
          "ALYVDSLFFL presented on HLA-A*02:01 and, simultaneously, CD3ε on T cells. The sponsor "
          "proposes to conduct Protocol " + PROTOCOL + ", a Phase 1, first-in-human, multi-centre, "
          "open-label, single-agent dose-escalation and expansion study in adults with locally advanced "
          "or metastatic solid tumours that are PRAME-positive and HLA-A*02:01-positive.",
          "12 Sections A.2, B.3")

    cited(doc,
          "The application contains the Form FDA 1571, this cover letter, an introductory statement and "
          "general investigational plan, the Investigator's Brochure, the clinical protocol, chemistry "
          "manufacturing and control information, and pharmacology and toxicology information. No "
          "clinical data are included; PMX-103 has not previously been administered to humans.",
          "12 Section B; 21 CFR 312.23(a)")

    cited(doc,
          "The proposed starting dose is 0.5 μg administered as a flat intravenous dose, selected using "
          "a modified minimum anticipated biological effect level approach. Mandatory step-up priming is "
          "employed, with 0.5 μg on Cycle 1 Day 1 and 5 μg on Cycle 1 Day 8 before the assigned cohort "
          "dose, together with protocol-specified premedication and cytokine release syndrome monitoring.",
          "12 Sections B.3, D.1; 10_Protocol §4.3")

    cited(doc,
          "The sponsor requests the Agency's review under 21 CFR Part 312. The sponsor understands that "
          "the study may not begin until 30 days after the date of receipt of this application unless "
          "the Agency notifies the sponsor earlier, and commits to the obligations of 21 CFR 312.50 "
          "through 312.70.",
          "21 CFR 312.40(b)")

    doc.add_paragraph()
    doc.add_paragraph("Questions regarding this application may be directed to the sponsor contact below.")
    doc.add_paragraph()
    for line in ["Sincerely,", "", "Ms. Laura Whitfield, JD, RAC",
                 "Vice President, Regulatory Affairs", SPONSOR, SPONSOR_ADDR,
                 "+1 410 555 0143"]:
        doc.add_paragraph(line)

    note(doc, "Signature block reproduced from the sponsor's draft. Signatory confirmation is GAP-02; "
              "the form must be signed by an authorised representative before submission.")


def _general_plan(doc):
    h(doc, 3, "1.20.1 Introductory Statement")
    cited(doc,
          f"{DRUG} is a recombinant humanised bispecific T-cell engager of approximately 55 kDa, "
          "constructed as a single-chain (scFv)₂ without an Fc domain and carrying a C-terminal "
          "hexahistidine C-tag. The domain arrangement is (VH–VL) anti-PRAME-TCRm — (G4S)₃ — (VL–VH) "
          "anti-CD3ε.",
          "12 Section A.2")
    cited(doc,
          "The molecule binds two targets simultaneously: the PRAME-derived nonamer ALYVDSLFFL, "
          "corresponding to PRAME residues 425–433 and presented on HLA-A*02:01, and CD3ε on T cells. "
          "Bridging the two brings polyclonal T cells into contact with tumour cells, producing T-cell "
          "activation, granzyme- and perforin-mediated lysis, release of interferon-γ, tumour necrosis "
          "factor-α and interleukin-2, and T-cell proliferation.",
          "12 Section A.2")
    cited(doc,
          "The drug product is a sterile solution for intravenous administration, supplied as 1.0 mg in "
          "1.1 mL in a single-use vial and stored at 2–8 °C protected from light.",
          "12 Section A.2; 07_CMC 3.2.P.1")
    cited(doc,
          "No proprietary name or International Nonproprietary Name has been assigned; an INN request is "
          "pending. The product is referred to throughout by its code, PMX-103.",
          "12 Section A.2 — GAP-03")

    h(doc, 3, "1.20.2 Rationale for the Drug and the Research")
    cited(doc,
          "PRAME is expressed in a range of solid tumours while being largely restricted in normal adult "
          "tissue, and its processed peptide is presented on HLA-A*02:01. That presentation makes an "
          "intracellular antigen accessible to an antibody-based therapeutic, which a conventional "
          "surface-binding antibody cannot reach. PMX-103 exploits this by pairing a T-cell receptor "
          "mimic arm against the peptide–MHC complex with a CD3ε-engaging arm.",
          "12 Section A.2; 08_IB §2")
    cited(doc,
          "Clinical precedent for PRAME-directed therapy in humans supports the target. The Pr20 T-cell "
          "receptor mimic antibody established that the ALYVDSLFFL/HLA-A*02:01 complex can be engaged "
          "therapeutically; IMA402, a PRAME-directed T-cell engaging receptor, and IMA203, a PRAME "
          "TCR-T cell product, have both entered clinical study.",
          "12 Section B.9; 08_IB §6")
    cited(doc,
          "In vitro, PMX-103 lyses PRAME-positive, HLA-A*02:01-positive A375-MA1 cells co-cultured with "
          "human peripheral blood mononuclear cells with a half-maximal effective concentration of "
          "0.62 ng/mL and a ten-percent effective concentration of 0.12 ng/mL.",
          "12 Section D.2 (PMX-103-IVT-002)")

    h(doc, 3, "1.20.3 Indication to be Studied")
    cited(doc,
          "The proposed indication is locally advanced or metastatic solid tumours that are "
          "PRAME-positive by central immunohistochemistry with an H-score of at least 80 and "
          "HLA-A*02:01-positive, in patients who have received at least one prior line of standard "
          "therapy including a PD-(L)1 inhibitor where indicated.",
          "12 Section A.2")

    h(doc, 3, "1.20.4 General Approach to Evaluating the Drug")
    cited(doc,
          "Development begins with a single first-in-human study combining dose escalation with "
          "disease-specific expansion. Escalation uses a modified toxicity probability interval design "
          "(mTPI-2) with an accelerated titration lead-in, a target dose-limiting toxicity rate of 30 "
          "percent and an equivalence interval of 0.25 to 0.35. Expansion at the recommended dose is "
          "intended to characterise activity in defined tumour types before any decision to proceed.",
          "12 Section B.3; 10_Protocol §§4, 9.3")
    cited(doc,
          "Dosing is once weekly by intravenous infusion in 21-day cycles. Mandatory step-up priming is "
          "used to mitigate cytokine release syndrome: 0.5 μg on Cycle 1 Day 1, 5 μg on Cycle 1 Day 8, "
          "and the assigned cohort dose from Cycle 2 Day 1 onward.",
          "12 Section B.3; 10_Protocol §4.3")
    cited(doc,
          "Premedication before each of the first doses comprises acetaminophen 650 mg orally, "
          "diphenhydramine 25 mg orally and dexamethasone 8 mg intravenously, given four to six hours "
          "before infusion.",
          "14 Invariants (Protocol §6.2)")

    h(doc, 3, "1.20.5 Kinds of Clinical Trials Planned for the First Year")
    cited(doc,
          "One study is planned in the first year following submission: Protocol " + PROTOCOL + ", the "
          "Phase 1 first-in-human study described above, comprising a dose-escalation part and a "
          "dose-expansion part. No other clinical study is planned in that period.",
          "12 Section B.3")
    cited(doc,
          "Dose escalation is expected to run approximately 24 months, and an individual subject may "
          "receive treatment for up to 24 months.",
          "12 Section B.3")

    h(doc, 3, "1.20.6 Estimated Number of Subjects")
    cited(doc,
          "Up to 50 subjects are planned: no more than 30 in the dose-escalation part and approximately "
          "20 in the expansion part, across four to six investigational sites.",
          "12 Section B.3")
    assume(doc,
           "Enrolment is not capped by a formal sample size calculation, because the study is a "
           "dose-finding investigation whose stopping behaviour is governed by the escalation design "
           "rather than by a power calculation.",
           "10_Protocol §9.3",
           assumption=(
               "Assumption: no formal power calculation underlies the 50-subject ceiling. This is normal "
               "for an mTPI-2 dose-finding study, and the input package gives operating characteristics "
               "rather than a sample size derivation. If the statistical analysis plan does contain a "
               "sample size justification, it should replace this paragraph."),
           anchor="not capped by a formal sample size calculation")

    h(doc, 3, "1.20.7 Anticipated Risks of Particular Severity or Seriousness")
    cited(doc,
          "The principal anticipated risk is cytokine release syndrome. In cynomolgus monkeys, single "
          "intravenous doses of 50 and 500 μg/kg produced fever, lethargy and a twofold rise in alanine "
          "aminotransferase, with T-cell redistribution within six hours. In HLA-A2.1 transgenic mice, "
          "interferon-γ and tumour necrosis factor-α peaked one to four hours after dosing and resolved "
          "by day three.",
          "12 Section D.4 (PMX-103-TOX-004, TOX-002)")
    cited(doc,
          "Neurological toxicity is anticipated as a class effect of T-cell engagers. Reversible "
          "neuroinflammation-like signs were observed in three of eight cynomolgus monkeys at 50 μg/kg "
          "in the repeat-dose study.",
          "12 Section D.4 (PMX-103-TOX-005)")
    cited(doc,
          "Hepatic enzyme elevation was reversible and confined to the highest doses tested; one death "
          "in eight animals occurred at 100 μg/kg in mice and was attributed to cytokine release.",
          "12 Section D.4 (PMX-103-TOX-003, TOX-006)")
    cited(doc,
          "Cardiovascular safety pharmacology showed no QTcF prolongation beyond 10 ms, with a maximum "
          "change of 9.2 ms and a transient 18 percent increase in heart rate at 50 μg/kg.",
          "12 Sections D.5, D.7 (PMX-103-SP-001)")
    cited(doc,
          "Anti-drug antibodies were detected in two of eight cynomolgus monkeys at day 28 of the "
          "repeat-dose study.",
          "12 Section D.4 (PMX-103-TOX-005)")
    cited(doc,
          "Risk mitigation in the protocol comprises the mandatory step-up priming schedule, "
          "premedication, inpatient observation windows, ASTCT 2019 grading for cytokine release "
          "syndrome and ASTCT ICE 2L v2 2024 grading for immune effector cell-associated neurotoxicity "
          "syndrome, together with protocol-defined stopping rules.",
          "12 Section B.3; 14 Invariants (Protocol §8.5)")

"""Build the authored PMX-103 IND document.

    python scripts/build_ind_document.py
"""
from __future__ import annotations

import sys
from pathlib import Path

SRC = Path(__file__).resolve().parents[1] / "src"
if str(SRC) not in sys.path:
    sys.path.insert(0, str(SRC))

from docx import Document                                    # noqa: E402
from docx.shared import Pt                                   # noqa: E402

from indkit.docgen import common                             # noqa: E402
from indkit.docgen.front_m1 import build_front, build_m1     # noqa: E402
from indkit.docgen.m2 import build_m2                        # noqa: E402
from indkit.docgen.m3_m5 import (build_appendices, build_m3, # noqa: E402
                                 build_m4, build_m5)
from indkit.pipeline.paths import REPORTS_DIR                # noqa: E402

OUT = REPORTS_DIR / "PMX-103_IND_Draft_IRIS.docx"


def main() -> int:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)

    build_front(doc)
    build_m1(doc)
    build_m2(doc)
    build_m3(doc)
    build_m4(doc)
    build_m5(doc)
    build_appendices(doc)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"saved {OUT}")
    print(f"  paragraphs {len(doc.paragraphs)} | tables {len(doc.tables)} | "
          f"comments {len(doc.comments)} | traceability rows {len(common.TRACE)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

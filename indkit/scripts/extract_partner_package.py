"""Turn the partner's PMX-103 input package into the two files the pipeline reads.

The package itself is not in version control (see .gitignore): the documents carry
the partner's own 21 CFR 312.130 confidentiality marking. What is committed is the
*output* of this script, so the mapping stays auditable without redistributing the
partner's files.

    cd indkit

Two files are written:

  data/source_cases/PMX103.json          the case the pipeline runs, in exactly the
                                         same shape as the ten fictional cases
  data/partner_reference/PMX103.json     the partner's own answers - GAPS log,
                                         section map, invariants, and the 1571 draft
                                         he hand-filled - kept separate because they
                                         are what we check ourselves against, never
                                         an input to generation

Everything written here is either copied verbatim from the package or assembled by
joining verbatim strings under a label. Nothing is paraphrased and nothing is
invented; a field the package does not answer is written as null so the pipeline
reports it as a gap.
"""
from __future__ import annotations

import json
import re
import sys
from pathlib import Path
from typing import Any

POC_DIR = Path(__file__).resolve().parent.parent
REPO_ROOT = POC_DIR.parent.parent
PACKAGE_DIR = REPO_ROOT / "inputs" / "partner-package" / "AI Inputs"

SOURCE_CASE_OUT = POC_DIR / "data" / "source_cases" / "PMX103.json"
REFERENCE_OUT = POC_DIR / "data" / "partner_reference" / "PMX103.json"

CASE_ID = "PMX103"

FILES = {
    "form_1571": "06_Form_FDA_1571_AutoFilled.docx",
    "cmc": "07_CMC_Dossier_PMX103_v3.docx",
    "ib": "08_Investigator_Brochure_PMX103_v3.docx",
    "ncd": "09_Nonclinical_Development_Plan_PMX103_v3.docx",
    "protocol": "10_Phase1_Protocol_PMX103_P1_001_v3.docx",
    "tox": "11_Toxicology_Dossier_PMX103_v3.docx",
    "brief": "12_IND_Drafting_Input_Brief_PMX103.docx",
    "instructions": "13_IND_Drafting_Instructions_PMX103.docx",
    "matrix": "14_IND_Section_Traceability_Matrix.xlsx",
    "schema": "15_AI_Agent_Input_Schema.json",
}


# --------------------------------------------------------------------------- io


def _require_package() -> None:
    if PACKAGE_DIR.is_dir():
        return
    print(f"error: partner package not found at {PACKAGE_DIR}", file=sys.stderr)
    print("This script only runs where the partner's input package is present.", file=sys.stderr)
    print("Its committed output (data/source_cases/PMX103.json) needs no re-run.", file=sys.stderr)
    raise SystemExit(2)


def _paragraphs(name: str) -> list[str]:
    import docx  # imported lazily so the pipeline itself never needs python-docx

    document = docx.Document(str(PACKAGE_DIR / FILES[name]))
    return [p.text for p in document.paragraphs]


def _schema() -> dict[str, Any]:
    return json.loads((PACKAGE_DIR / FILES["schema"]).read_text(encoding="utf-8"))


def _sheets() -> dict[str, list[list[Any]]]:
    import openpyxl  # lazily, same reason

    workbook = openpyxl.load_workbook(str(PACKAGE_DIR / FILES["matrix"]))
    out: dict[str, list[list[Any]]] = {}
    for sheet in workbook:
        rows = []
        for row in sheet.iter_rows(values_only=True):
            if all(cell is None for cell in row):
                continue
            rows.append(["" if cell is None else str(cell).strip() for cell in row])
        out[sheet.title] = rows
    return out


def _rows_as_dicts(rows: list[list[Any]], keys: list[str]) -> list[dict[str, str]]:
    """Header row is dropped; the caller supplies the key names it wants."""
    return [dict(zip(keys, row)) for row in rows[1:]]


# ------------------------------------------------------- his Form 1571 draft


NUMBERED = re.compile(r"^\s*(\d+)\.\s*([^:]+):\s*(.*)$")
UNNUMBERED = re.compile(r"^\s*([^:]+):\s*(.*)$")


def parse_form_1571_draft() -> list[dict[str, Any]]:
    """Read 06_Form_FDA_1571_AutoFilled.docx into `label: value` entries.

    His draft is a flat list of `n. Label: value` lines, plus two lines above the
    numbering (IND number, serial number) that carry no number at all. Both shapes
    are captured; the numbering is recorded exactly as he wrote it, because the
    numbering itself is one of the findings.
    """
    entries: list[dict[str, Any]] = []
    for line in _paragraphs("form_1571"):
        text = line.strip()
        if not text or text.startswith("Form FDA 1571"):
            continue
        match = NUMBERED.match(text)
        if match:
            entries.append(
                {
                    "draft_number": match.group(1),
                    "label": match.group(2).strip(),
                    "value": match.group(3).strip(),
                }
            )
            continue
        match = UNNUMBERED.match(text)
        if match:
            entries.append(
                {
                    "draft_number": None,
                    "label": match.group(1).strip(),
                    "value": match.group(2).strip(),
                }
            )
    return entries


def field_from_draft(entries: list[dict[str, Any]], label_fragment: str) -> str | None:
    for entry in entries:
        if label_fragment.lower() in entry["label"].lower():
            return entry["value"]
    return None


# ------------------------------------------------------- the brief's B.3 clauses

# B.3 states the general investigational plan as a single semicolon-separated list
# of `label (value)` clauses. Splitting on that structure keeps every value verbatim.
CLAUSE = re.compile(r"^(?P<label>[^()]+?)\s*\((?P<value>.+)\)$")


def general_plan_clauses() -> dict[str, str]:
    for text in _paragraphs("brief"):
        marker = "General investigational plan:"
        if marker not in text:
            continue
        tail = text.split(marker, 1)[1].strip().rstrip(".")
        clauses: dict[str, str] = {}
        # Split on top-level semicolons only; the values themselves contain commas
        # and brackets but no semicolons.
        for chunk in tail.split(";"):
            match = CLAUSE.match(chunk.strip())
            if match:
                clauses[match.group("label").strip().lower()] = match.group("value").strip()
        if clauses:
            return clauses
    raise SystemExit("error: could not locate the B.3 general investigational plan clauses")


def anticipated_risks() -> str:
    """IB section 7.2 states the anticipated risks in one paragraph. Take it whole."""
    for text in _paragraphs("ib"):
        stripped = text.strip()
        if stripped.startswith("The anticipated risks of PMX-103 in the clinical study are:"):
            return stripped
    raise SystemExit("error: could not locate the IB 7.2 anticipated-risks paragraph")


def brief_table_value(table_index: int, field_name: str) -> str | None:
    import docx

    document = docx.Document(str(PACKAGE_DIR / FILES["brief"]))
    table = document.tables[table_index]
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        if cells and cells[0].lower() == field_name.lower():
            return cells[1]
    return None


def joined(*parts: tuple[str, str | None]) -> str | None:
    """Join `label: verbatim value` blocks. No connective prose is written."""
    kept = [f"{label}: {value}" for label, value in parts if value]
    return "\n\n".join(kept) if kept else None


# --------------------------------------------------------------- the source case


def build_source_case() -> dict[str, Any]:
    schema = _schema()
    program = schema["program"]
    draft = parse_form_1571_draft()
    clauses = general_plan_clauses()

    address = program["sponsor_addr"]  # "1515 Hollins Street, Suite 400, Baltimore, MD 21202"
    street, suite, city, state_zip = [part.strip() for part in address.split(",")]
    state, postal = state_zip.split()

    # The signatory named in the brief's A.1 table, cross-checked against authors[].
    signatory = next(a for a in schema["authors"] if a["name"].startswith("Ms. Laura Whitfield"))

    drug_line = field_from_draft(draft, "Drug Name")  # "PMX-103 (Generic: ... BiTE)"
    generic = None
    if drug_line and "Generic:" in drug_line:
        generic = drug_line.split("Generic:", 1)[1].strip().rstrip(")").strip()

    records = [
        {
            "record_id": "sponsor_profile",
            "record_type": "sponsor_profile",
            "title": "Sponsor identity from the partner input schema",
            "source_type": "partner_supplied_structured_data",
            "source_document": f"{FILES['schema']} (program), {FILES['form_1571']} (item 4)",
            "fields": {
                "sponsor.legal_name": program["sponsor"],
                "sponsor.address_line_1": street,
                "sponsor.address_line_2": suite,
                "sponsor.city": city,
                "sponsor.state": state,
                "sponsor.postal_code": postal,
                "sponsor.country": "USA",
                "sponsor.contact_name": signatory["name"],
                "sponsor.contact_title": signatory["title"],
                "sponsor.phone": field_from_draft(draft, "Telephone"),
                # Not answered anywhere in the package. Left null on purpose.
                "sponsor.email": None,
            },
        },
        {
            "record_id": "product_brief",
            "record_type": "investigational_product_brief",
            "title": "Investigational product identity",
            "source_type": "partner_supplied_structured_data",
            "source_document": f"{FILES['schema']} (program), {FILES['brief']} (A.2)",
            "fields": {
                "product.code_name": program["drug_code"],
                "product.generic_name": generic,
                "product.dosage_form": brief_table_value(3, "Dosage form / strength"),
                "product.route": "IV",
                "product.indication": program["indication"],
            },
        },
        {
            "record_id": "protocol_synopsis",
            "record_type": "clinical_protocol_synopsis",
            "title": "Protocol PMX-103-P1-001 document control block",
            "source_type": "partner_supplied_structured_data",
            "source_document": FILES["protocol"],
            "fields": {
                "protocol.protocol_number": "PMX-103-P1-001",
                "protocol.title": "Phase 1 Clinical Study Protocol for PMX-103 (PMX-103-P1-001)",
                "protocol.phase": "Phase 1",
                "protocol.version": "3.0",
                "protocol.protocol_date": "2026-08-22",
            },
        },
        {
            "record_id": "submission_intake",
            "record_type": "regulatory_submission_intake",
            "title": "Submission intake",
            "source_type": "partner_supplied_structured_data",
            "source_document": f"{FILES['schema']} (program.submit_date), {FILES['form_1571']}",
            "fields": {
                "submission.submission_type": "Initial IND",
                "submission.submission_date": "2026-08-21",
                "submission.serial_number": "0000",
                # His draft carries "100,001 (Sponsor reservation)". A reservation is
                # not an assigned number, so nothing is carried forward. GAP-01.
                "submission.ind_number": None,
            },
        },
        {
            "record_id": "investigator_profile",
            "record_type": "principal_investigator_profile",
            "title": "Principal investigator profile - not supplied by the package",
            "source_type": "partner_supplied_structured_data",
            "source_document": "none; sites and investigators are not selected (GAP-04)",
            "fields": {
                "investigator.name": None,
                "investigator.institution": None,
                "investigator.address": None,
                "investigator.phone": None,
                "investigator.email": None,
            },
        },
        {
            "record_id": "investigational_plan",
            "record_type": "general_investigational_plan",
            "title": "General investigational plan assembled from the brief B.3 clauses",
            "source_type": "partner_supplied_structured_data",
            "source_document": (
                f"{FILES['brief']} (B.3), {FILES['schema']} (program), {FILES['ib']} (7.2)"
            ),
            "fields": {
                "plan.rationale": joined(
                    ("Mechanism of action", program["moa"]),
                    ("Class precedent", program["key_prev"]),
                    ("Starting-dose rationale", program["starting_dose_rationale"]),
                ),
                "plan.general_approach": joined(
                    ("Dose-escalation approach", clauses.get("dose-escalation approach")),
                    ("Dose levels", program["dose_design"]),
                    ("Step-up priming", program["step_up"]),
                    ("Premedication", program["premed"]),
                ),
                "plan.first_year_scope": joined(
                    ("Study design", clauses.get("study design")),
                    ("Number of sites", clauses.get("number of sites")),
                    ("Planned duration", clauses.get("planned duration")),
                ),
                "plan.estimated_enrollment": clauses.get("number of subjects"),
                "plan.anticipated_risks": anticipated_risks(),
            },
            "planned_studies": [
                {
                    "study_id": "PMX-103-P1-001",
                    "title": "Phase 1 Clinical Study Protocol for PMX-103 (PMX-103-P1-001)",
                    "objectives": None,
                    "design": clauses.get("study design"),
                    "sample_size": clauses.get("number of subjects"),
                    "population": program["indication"],
                    "parameters": None,
                    "status": "Protocol v3.0 dated 22 August 2026. Enrolment has not started.",
                }
            ],
        },
    ]

    for record in records:
        studies = record.get("planned_studies")
        if studies:
            record["planned_studies"] = [
                {key: value for key, value in study.items() if value is not None}
                for study in studies
            ]

    return {
        "case_id": CASE_ID,
        "case_label": "PMX103 - PMX-103 partner-supplied Initial IND input",
        "scenario_type": "missing",
        "origin": "partner_supplied",
        "display_name": "PMX-103",
        "partner_package": {
            "program": "PMX-103 (Indela Therapeutics, Inc.)",
            "supplied_by": "clinical partner",
            "package_date": _schema()["generated"],
            "documents": sorted(FILES.values()),
            "confidentiality": "Partner documents carry a 21 CFR 312.130 marking and are not committed.",
        },
        "scope_boundary": scope_boundary(schema),
        "source_records": records,
        "notes": [
            "Real-shaped input supplied by the clinical partner, mapped into the same canonical "
            "schema the ten fictional cases use. Values are copied verbatim from the package.",
            "sponsor.email and every investigator.* field are unanswered by the package. They are "
            "left empty so the pipeline reports them rather than filling them in.",
            "submission.ind_number is left empty: the partner's 1571 draft carries a reserved "
            "number (100,001), which is not an assigned IND number.",
        ],
    }


def scope_boundary(schema: dict[str, Any]) -> dict[str, Any]:
    """State plainly how much of the package Module 1 can actually consume.

    This is counted from the package, not asserted: the numbers below are the
    lengths of the schema's own tables.
    """
    return {
        "headline": "Module 1 uses a narrow slice of this package. That is the scope, not a shortfall.",
        "consumed": (
            "Sponsor identity, product identity, protocol identity, submission intake and the "
            "1.20 general investigational plan - about 35 scalar fields."
        ),
        "not_consumed": [
            {
                "what": f"{len(schema['dose_levels'])} dose levels and the mTPI-2 escalation design",
                "belongs_to": "Module 5 (protocol) / Module 2.7",
            },
            {
                "what": f"{len(schema['toxicology_studies'])} toxicology studies and "
                f"{len(schema['safety_pharmacology'])} safety-pharmacology studies",
                "belongs_to": "Module 4, summarised in Module 2.6",
            },
            {
                "what": f"{len(schema['cmc_release_spec'])} CMC release specifications, lots and stability",
                "belongs_to": "Module 3, summarised in Module 2.3",
            },
            {
                "what": f"{len(schema['in_vitro_potency'])} in-vitro potency assays and the PK projection",
                "belongs_to": "Module 4 / Module 2.6",
            },
        ],
        "position": (
            "Modules 2 to 5 are not built and are not planned for this demonstration. The partner's "
            "drafting instructions ask for a full M1-M5 dossier; this responds to Module 1 only."
        ),
    }


# ------------------------------------------------------------ partner reference


def brief_module_map() -> list[dict[str, str]]:
    """Section C of the brief: the eCTD module each source dossier belongs to.

    Kept because it is the partner's second, independent statement of the same
    mapping the traceability matrix makes - and the two do not agree.
    """
    import docx

    document = docx.Document(str(PACKAGE_DIR / FILES["brief"]))
    for table in document.tables:
        header = [cell.text.strip().lower() for cell in table.rows[0].cells]
        if header[:1] == ["ectd module"]:
            return [
                {
                    "module": row.cells[0].text.strip(),
                    "content": row.cells[1].text.strip(),
                    "source_documents": row.cells[2].text.strip(),
                }
                for row in table.rows[1:]
            ]
    raise SystemExit("error: could not locate the brief's Section C module map")


def build_partner_reference() -> dict[str, Any]:
    sheets = _sheets()
    schema = _schema()

    return {
        "case_id": CASE_ID,
        "package_date": schema["generated"],
        "package_files": sorted(p.name for p in PACKAGE_DIR.iterdir() if p.is_file()),
        "source_files": {
            "gaps_log": f"{FILES['matrix']} > GAPS Log",
            "section_map": f"{FILES['matrix']} > Section Map",
            "invariants": f"{FILES['matrix']} > Invariants",
            "form_1571_draft": FILES["form_1571"],
            "module_map": f"{FILES['brief']} > Section C",
            "instructions": FILES["instructions"],
        },
        "gaps_log": _rows_as_dicts(
            sheets["GAPS Log"], ["id", "item", "owner", "source_needed", "status"]
        ),
        "section_map": _rows_as_dicts(
            sheets["Section Map"],
            ["section", "ectd_location", "required_content", "source_document", "source_file", "status"],
        ),
        "invariants": _rows_as_dicts(
            sheets["Invariants"], ["parameter", "value", "unit", "source"]
        ),
        "module_map": brief_module_map(),
        "form_1571_draft": parse_form_1571_draft(),
    }


def main() -> int:
    _require_package()

    source_case = build_source_case()
    reference = build_partner_reference()

    SOURCE_CASE_OUT.parent.mkdir(parents=True, exist_ok=True)
    REFERENCE_OUT.parent.mkdir(parents=True, exist_ok=True)
    SOURCE_CASE_OUT.write_text(
        json.dumps(source_case, indent=2, ensure_ascii=False) + "\n", encoding="utf-8"
    )
    REFERENCE_OUT.write_text(
        json.dumps(reference, indent=2, ensure_ascii=False) + "\n", encoding="utf-8"
    )

    supplied = sum(
        1
        for record in source_case["source_records"]
        for value in record["fields"].values()
        if value not in (None, "")
    )
    blank = sum(
        1
        for record in source_case["source_records"]
        for value in record["fields"].values()
        if value in (None, "")
    )
    print(f"wrote {SOURCE_CASE_OUT.relative_to(POC_DIR)}  ({supplied} fields supplied, {blank} left empty)")
    print(
        f"wrote {REFERENCE_OUT.relative_to(POC_DIR)}  "
        f"({len(reference['gaps_log'])} gaps, {len(reference['invariants'])} invariants, "
        f"{len(reference['form_1571_draft'])} 1571 draft lines)"
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

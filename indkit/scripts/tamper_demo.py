"""Change one number in one document, re-run the scan, and see what comes back.

A report where all 23 invariants read the same is the right answer and a bad demo:
nobody can tell whether the scan looked. This makes the scan fail on purpose.

    cd indkit
    python build/tamper_demo.py --show                     what can be tampered with
    python build/tamper_demo.py --parameter noael_cyno     break it and re-scan
    python build/tamper_demo.py --clean                    delete the working copy

The partner's own files are never touched. The whole package is copied into
outputs/tamper/ first, the copy is edited, and the scan is pointed at the copy.
"""
from __future__ import annotations

import argparse
import json
import shutil
import sys
from pathlib import Path

POC_DIR = Path(__file__).resolve().parent.parent
REPO_ROOT = POC_DIR.parent.parent
PACKAGE_DIR = REPO_ROOT / "inputs" / "partner-package" / "AI Inputs"
WORK_DIR = POC_DIR / "outputs" / "tamper" / "AI Inputs"
REPORT = POC_DIR / "outputs" / "tamper" / "invariant_scan_tampered.json"
BASELINE = POC_DIR / "outputs" / "generated" / "PMX103" / "invariant_scan.json"

sys.path.insert(0, str(Path(__file__).resolve().parent))

# Which document to edit for a given parameter, and what to look for. Each entry names
# a phrase that appears verbatim in that document.
TARGETS: dict[str, dict[str, str]] = {
    "noael_cyno": {
        "document": "08_Investigator_Brochure_PMX103_v3.docx",
        "find": "NOAEL",
        "from": "NOAEL was 5 micrograms/kg",
        "to": "NOAEL was 8 micrograms/kg",
        "what": "the cynomolgus NOAEL, in the Investigator's Brochure",
    },
    "starting_dose_ug_flat": {
        "document": "08_Investigator_Brochure_PMX103_v3.docx",
        "find": "starting dose",
        "from": "starting dose of 0.5 microgram flat",
        "to": "starting dose of 0.8 microgram flat",
        "what": "the proposed starting dose, in the Investigator's Brochure",
    },
    "in_vitro_EC50": {
        "document": "09_Nonclinical_Development_Plan_PMX103_v3.docx",
        "find": "EC50",
        "from": "EC50 of 0.62 ng/mL",
        "to": "EC50 of 0.72 ng/mL",
        "what": "the in-vitro EC50, in the Nonclinical Development Plan",
    },
}


def sync_working_copy() -> None:
    if WORK_DIR.exists():
        shutil.rmtree(WORK_DIR)
    WORK_DIR.parent.mkdir(parents=True, exist_ok=True)
    shutil.copytree(PACKAGE_DIR, WORK_DIR)


def tamper(target: dict[str, str]) -> int:
    """Rewrite every paragraph in the working copy that carries the target phrase.

    Runs inside a paragraph can split a phrase anywhere, so the replacement is done at
    paragraph level: the new text goes into the first run and the rest are emptied.
    Formatting inside the paragraph is lost, which does not matter in a throwaway copy.
    """
    import docx

    path = WORK_DIR / target["document"]
    document = docx.Document(str(path))
    changed = 0

    def rewrite(paragraph) -> bool:
        if target["from"] not in paragraph.text:
            return False
        if target["find"].lower() not in paragraph.text.lower():
            return False
        updated = paragraph.text.replace(target["from"], target["to"])
        for index, run in enumerate(paragraph.runs):
            run.text = updated if index == 0 else ""
        if not paragraph.runs:
            paragraph.add_run(updated)
        return True

    for paragraph in document.paragraphs:
        changed += rewrite(paragraph)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    changed += rewrite(paragraph)

    document.save(str(path))
    return changed


def baseline_status(parameter: str) -> str:
    if not BASELINE.exists():
        return "unknown (no baseline report yet)"
    report = json.loads(BASELINE.read_text(encoding="utf-8"))
    for row in report["invariants"]:
        if row["parameter"] == parameter:
            return f"{row['status']} across {len(row.get('documents', []))} document(s)"
    return "unknown"


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__.splitlines()[0])
    parser.add_argument("--parameter", choices=sorted(TARGETS), help="Which invariant to break.")
    parser.add_argument("--show", action="store_true", help="List what can be tampered with.")
    parser.add_argument("--clean", action="store_true", help="Delete the working copy and stop.")
    args = parser.parse_args()

    if args.clean:
        if WORK_DIR.parent.exists():
            shutil.rmtree(WORK_DIR.parent)
            print(f"removed {WORK_DIR.parent}")
        else:
            print("nothing to remove")
        return 0

    if args.show or not args.parameter:
        print("Available edits (the partner's own files are never modified):\n")
        for parameter, target in sorted(TARGETS.items()):
            print(f"  {parameter}")
            print(f"    {target['what']}")
            print(f"    {target['from']}  ->  {target['to']}")
            print(f"    baseline: {baseline_status(parameter)}\n")
        return 0

    if not PACKAGE_DIR.is_dir():
        print(f"error: partner package not found at {PACKAGE_DIR}", file=sys.stderr)
        return 2

    target = TARGETS[args.parameter]
    print(f"baseline for {args.parameter}: {baseline_status(args.parameter)}")
    print(f"copying the package to {WORK_DIR} ...")
    sync_working_copy()

    changed = tamper(target)
    if not changed:
        print(f"error: phrase “{target['from']}” not found in {target['document']}", file=sys.stderr)
        return 1
    print(f"changed {changed} paragraph(s) in {target['document']}: "
          f"{target['from']} -> {target['to']}\n")

    from scan_invariants import main as scan

    code = scan([
        "--package-dir", str(WORK_DIR),
        "--out", str(REPORT),
        "--only", args.parameter,
    ])
    if code != 0:
        return code

    report = json.loads(REPORT.read_text(encoding="utf-8"))
    row = next(r for r in report["invariants"] if r["parameter"] == args.parameter)

    print(f"\n{'=' * 72}")
    print(f"{args.parameter}: {row['status']}")
    print(f"declared {row['declared_value']} {row['declared_unit']} ({row['declared_source']})")
    print(f"{'=' * 72}")
    for assertion in row["assertions"]:
        mark = "  ok " if assertion["agrees"] else "  >> "
        print(f"{mark}{assertion['document']} @ char {assertion['offset']}")
        print(f"      asserts {assertion['asserted_value']} {assertion['asserted_unit'] or ''}".rstrip())
        if not assertion["agrees"]:
            print(f"      {assertion['why']}")
            print(f"      “{assertion['sentence'][:200]}”")
    print(f"\nfull report: {REPORT}")
    print("run with --clean to remove the working copy")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

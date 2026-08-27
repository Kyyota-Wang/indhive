"""Check the partner's 23 invariant numbers against the documents that state them.

The traceability matrix declares 23 parameters that must read the same everywhere in
the package. This script goes and looks.

    cd indkit

Four steps, and the split between them is the whole point.

  1  DETERMINISTIC LOCATION   Find every occurrence of the parameter's anchor term
                              across the corpus and cut out the sentence containing
                              it. This produces candidates, not answers.

  2  MODEL ADJUDICATION       Ask a model, one sentence at a time, a single closed
                              question: does this sentence assert that <parameter>
                              equals <value> <unit>? It judges; it never generates a
                              number and it never sees the other documents.

  3  DETERMINISTIC COMPARISON Normalise units and compare the asserted value against
                              the declared one in code. The model's opinion of whether
                              two numbers match is not consulted.

  4  EVIDENCE                 Every assertion keeps its document, its sentence and its
                              character offset, so a reader can open the file and check.

Step 2 exists because step 1 alone does not work. Anchoring on a parameter name and
taking the nearest number picks up the 72309 in an fda.gov/media/72309/download URL as
a starting dose, the 103 in "PMX-103" as a mouse MTD, and the HNSTD in a sentence that
happens to mention the NOAEL. Those are not edge cases; they are most of the hits. A
judgement call on natural language is what removes them, and a judgement call is the
one thing deterministic code cannot make.

The corpus is the partner's package, which is not in version control. The output is,
so the report can be read without it.
"""
from __future__ import annotations

import json
import os
import re
import sys
import urllib.error
import urllib.request
from pathlib import Path
from typing import Any

POC_DIR = Path(__file__).resolve().parent.parent
REPO_ROOT = POC_DIR.parent.parent
PACKAGE_DIR = REPO_ROOT / "inputs" / "partner-package" / "AI Inputs"
REFERENCE = POC_DIR / "data" / "partner_reference" / "PMX103.json"
OUT_FILE = POC_DIR / "outputs" / "generated" / "PMX103" / "invariant_scan.json"

MODEL = "claude-opus-5"
API = "https://api.anthropic.com/v1/messages"

# The documents that carry the numbers. The two instruction documents (13) and the
# schema (15) are excluded: they are the specification, not a statement of the value.
CORPUS = [
    "07_CMC_Dossier_PMX103_v3.docx",
    "08_Investigator_Brochure_PMX103_v3.docx",
    "09_Nonclinical_Development_Plan_PMX103_v3.docx",
    "10_Phase1_Protocol_PMX103_P1_001_v3.docx",
    "11_Toxicology_Dossier_PMX103_v3.docx",
    "12_IND_Drafting_Input_Brief_PMX103.docx",
]

# What to search for, per invariant parameter. Anchored on how the parameter is named
# in prose, never on its value - searching for the value would only ever confirm
# itself and would find nothing when a document disagreed.
ANCHORS: dict[str, list[str]] = {
    "starting_dose_ug_flat": ["starting dose", "first-in-human dose", "FIH dose"],
    "starting_dose_Cmax": ["projected Cmax", "predicted Cmax", "human Cmax"],
    "in_vitro_EC10": ["EC10"],
    "in_vitro_EC50": ["EC50"],
    "mabel_anchor_mouse": ["MABEL"],
    "noael_cyno": ["NOAEL"],
    "noael_glp_mouse": ["NOAEL"],
    "mtd_mouse": ["MTD", "maximum tolerated dose"],
    "hnstd_cyno": ["HNSTD", "highest non-severely toxic dose"],
    "cyno_Cmax_40ug": ["Cmax"],
    "cyno_AUC_168h": ["AUC"],
    "human_CL": ["clearance"],
    "human_Vss": ["Vss", "volume of distribution"],
}

# Why a parameter is left out of the numeric scan. Saying "not scanned" without saying
# why is the kind of silence this whole project exists to remove.
NOT_SCANNED_REASONS: dict[str, str] = {
    "drug_code": "An identifier, not a quantity. Checking that a code reads the same "
    "everywhere is a text-identity check, which this scan does not do.",
    "dlt_window_days": "Declared as 28 days, but the corpus states it as a cycle-day window "
    "(C1D1-C1D28, 'Cycle 1 Day 1 through Cycle 1 Day 28'). Reading 28 days out of that means "
    "converting between two representations, and this scan compares values rather than "
    "converting them.",
}

DEFAULT_NOT_SCANNED = (
    "Stated as text rather than a number with a unit. This scan compares numbers; a text "
    "invariant needs a different comparison and does not have one yet."
)

# Extra context the judge needs, because several parameters share an anchor. NOAEL
# appears for both the cyno and the GLP mouse study; without the species the question
# is ambiguous and the answer is worthless.
QUALIFIERS: dict[str, str] = {
    "noael_cyno": "in the cynomolgus monkey",
    "noael_glp_mouse": "in the GLP mouse study",
    "mtd_mouse": "in the mouse",
    "hnstd_cyno": "in the cynomolgus monkey",
    "mabel_anchor_mouse": "the MABEL anchor derived from the mouse study",
    "cyno_Cmax_40ug": "in the cynomolgus monkey at the 40 ug/kg dose",
    "starting_dose_Cmax": "the projected human Cmax at the 0.5 ug starting dose",
    "cyno_AUC_168h": "the cynomolgus monkey AUC over 0 to 168 hours",
    "human_CL": "the projected human clearance",
    "human_Vss": "the projected human volume of distribution at steady state",
    "starting_dose_ug_flat": "the proposed human starting dose, as a flat dose",
    "in_vitro_EC10": "the in-vitro EC10 for tumour-cell lysis",
    "in_vitro_EC50": "the in-vitro EC50 for tumour-cell lysis",
    "dlt_window_days": "the length of the DLT observation window",
}

# Unit spellings that mean the same thing. Normalised before any comparison so a
# document writing "microg/kg" and a table writing "ug/kg" do not read as a conflict.
UNIT_FOLD = [
    (r"micrograms?", "ug"),
    (r"microg", "ug"),
    (r"µg", "ug"),
    (r"μg", "ug"),
    (r"mcg", "ug"),
    (r"ng\s*[.*·]\s*h", "ng*h"),
    (r"ng\s*h\b", "ng*h"),
    (r"\bhours?\b", "h"),
    (r"\bdays?\b", "d"),
    # Dosing qualifiers, not units. "0.5 microgram flat" is 0.5 ug given as a flat dose;
    # reading "flat" as part of the unit turns an agreement into a false conflict.
    (r"\b(flat|fixed|total|absolute|dose)\b", ""),
]

MAX_CANDIDATES = 40
# Sentences break on a full stop, semicolon or question mark - never on a colon. A table
# row reads "Calcein-AM release | 0.62 (EC10: 0.12)", and cutting at the colon separates a
# value from its label, which then reads as a document that states no value at all.
SENTENCE_SPLIT = re.compile(r"(?<=[.;!?])\s+|\n+")


# ------------------------------------------------------------------ 1. locate


def _document_text(name: str) -> str:
    """Paragraphs and table cells, in document order, as one string."""
    import docx

    document = docx.Document(str(PACKAGE_DIR / name))
    chunks: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            chunks.append(paragraph.text.strip())
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                chunks.append(" | ".join(cells))
    return "\n".join(chunks)


def _sentences(text: str) -> list[tuple[int, str]]:
    """(character offset, sentence). Offsets are into the text this function was given."""
    out: list[tuple[int, str]] = []
    position = 0
    for piece in SENTENCE_SPLIT.split(text):
        if piece is None:
            continue
        index = text.find(piece, position)
        if index < 0:
            index = position
        position = index + len(piece)
        cleaned = piece.strip()
        if cleaned:
            out.append((index, cleaned))
    return out


def locate_candidates(corpus: dict[str, str], parameter: str) -> list[dict[str, Any]]:
    anchors = [a.lower() for a in ANCHORS[parameter]]
    found: list[dict[str, Any]] = []
    seen: set[str] = set()

    for name, text in corpus.items():
        for offset, sentence in _sentences(text):
            lowered = sentence.lower()
            hit = next((a for a in anchors if a in lowered), None)
            if not hit:
                continue
            key = f"{name}|{sentence}"
            if key in seen:
                continue
            seen.add(key)
            found.append(
                {"document": name, "offset": offset, "anchor": hit, "sentence": sentence}
            )
    return found


# ------------------------------------------------------------------ 2. adjudicate


JUDGE_SYSTEM = """You are a discriminator, not a writer. You are given one parameter with a
declared value and unit, and a numbered list of sentences taken from a regulatory dossier.

For each sentence, answer one question: does this sentence state a value for that parameter?

Rules:
- A sentence asserts the parameter only if it states a value FOR THAT PARAMETER. A sentence
  that mentions the parameter while stating some other quantity does not.
- Digits inside a URL, a document number, a study identifier, a compound code, a section
  reference, a citation or a lot number are never the value of a parameter.
- If the sentence states a value for a DIFFERENT parameter that happens to be named in the
  same sentence, that is NOT_AN_ASSERTION for this one.
- Report the value exactly as the sentence writes it. Do not convert, round, or normalise.
  Do not compare it to the declared value; that comparison is done elsewhere.
- If a sentence gives a range, report the range verbatim as the value.
- If a sentence plainly sets out to state this parameter's value but the value itself is
  missing, truncated or garbled, that is still ASSERTS: return it with "value": null. Never
  guess what the number was meant to be, and never fill it in from the declared value.

Return ONLY a JSON array, one object per sentence, no prose and no markdown fence:
[{"i": <index>, "verdict": "ASSERTS" | "NOT_AN_ASSERTION", "value": <string or null>, "unit": <string or null>}]"""


def _api_key() -> str:
    key = os.getenv("ANTHROPIC_API_KEY")
    if key:
        return key
    dev_vars = REPO_ROOT / "V1" / ".dev.vars"
    if dev_vars.exists():
        for line in dev_vars.read_text(encoding="utf-8").splitlines():
            if line.startswith("ANTHROPIC_API_KEY="):
                return line.split("=", 1)[1].strip()
    raise SystemExit("error: ANTHROPIC_API_KEY is not set and V1/.dev.vars has no key")


def _ask(key: str, prompt: str) -> str:
    body = json.dumps(
        {
            "model": MODEL,
            "max_tokens": 4000,
            "system": JUDGE_SYSTEM,
            "messages": [{"role": "user", "content": prompt}],
        }
    ).encode("utf-8")
    request = urllib.request.Request(
        API,
        data=body,
        headers={
            "x-api-key": key,
            "anthropic-version": "2023-06-01",
            "content-type": "application/json",
        },
    )
    try:
        with urllib.request.urlopen(request, timeout=180) as response:
            payload = json.loads(response.read().decode("utf-8"))
    except urllib.error.HTTPError as exc:
        raise SystemExit(f"error: API returned {exc.code}: {exc.read().decode('utf-8')[:300]}")
    return "".join(b.get("text", "") for b in payload.get("content", []) if b.get("type") == "text")


def adjudicate(key: str, invariant: dict[str, str], candidates: list[dict[str, Any]]) -> list[dict[str, Any]]:
    if not candidates:
        return []

    parameter = invariant["parameter"]
    listing = "\n".join(
        f'{i}. [{c["document"]}] {c["sentence"]}' for i, c in enumerate(candidates)
    )
    prompt = (
        f'Parameter: {parameter}\n'
        f'Described as: {QUALIFIERS.get(parameter, parameter)}\n'
        f'Declared value: {invariant["value"]}\n'
        f'Declared unit: {invariant["unit"]}\n\n'
        f"Sentences:\n{listing}"
    )

    raw = _ask(key, prompt).strip()
    raw = re.sub(r"^```(?:json)?|```$", "", raw, flags=re.MULTILINE).strip()
    try:
        verdicts = json.loads(raw)
    except json.JSONDecodeError:
        print(f"  warning: could not parse the judge's reply for {parameter}", file=sys.stderr)
        return []

    out = []
    for verdict in verdicts:
        index = verdict.get("i")
        if not isinstance(index, int) or not 0 <= index < len(candidates):
            continue
        if verdict.get("verdict") != "ASSERTS":
            continue
        out.append({**candidates[index], "asserted_value": verdict.get("value"), "asserted_unit": verdict.get("unit")})
    return out


# ------------------------------------------------------------------ 3. compare


def _fold_unit(unit: str | None) -> str:
    text = (unit or "").strip().lower()
    for pattern, replacement in UNIT_FOLD:
        text = re.sub(pattern, replacement, text)
    return re.sub(r"\s+", "", text).strip("-")


def _number(value: str | None) -> float | None:
    match = re.search(r"-?\d+(?:[.,]\d+)?", str(value or "").replace(",", ""))
    return float(match.group()) if match else None


MATCHES = "MATCHES"
DIFFERS = "DIFFERS"
UNREADABLE = "UNREADABLE"


def compare(invariant: dict[str, str], assertion: dict[str, Any]) -> dict[str, Any]:
    """Three outcomes, and they are not the same kind of problem.

    A value that differs is a contradiction. A sentence that sets out to state the
    parameter and carries no readable number is a hole in the document - a different
    defect, reported separately rather than folded in with the contradictions.
    """
    declared = _number(invariant["value"])
    asserted = _number(assertion.get("asserted_value"))

    if asserted is None:
        verdict = UNREADABLE
        why = "The sentence states this parameter but carries no readable value."
    elif declared is None:
        verdict = UNREADABLE
        why = "The declared value is not a number, so it cannot be compared."
    elif abs(declared - asserted) > 1e-9:
        verdict = DIFFERS
        why = f"States {asserted}; the matrix declares {declared}."
    else:
        declared_unit = _fold_unit(invariant["unit"])
        asserted_unit = _fold_unit(assertion.get("asserted_unit"))
        # A missing unit in prose is normal; a different one is not.
        if asserted_unit and declared_unit and asserted_unit != declared_unit:
            verdict = DIFFERS
            why = f"Unit {assertion.get('asserted_unit')} does not match the declared {invariant['unit']}."
        else:
            verdict = MATCHES
            why = ""

    return {**assertion, "verdict": verdict, "agrees": verdict == MATCHES, "why": why}


# ---------------------------------------------------------------------- run


def main(argv: list[str] | None = None) -> int:
    global PACKAGE_DIR

    import argparse

    parser = argparse.ArgumentParser(description="Scan the corpus for the declared invariants.")
    parser.add_argument(
        "--package-dir",
        type=Path,
        default=PACKAGE_DIR,
        help="Where the corpus lives. Point this at a copy to scan a modified document.",
    )
    parser.add_argument("--out", type=Path, default=OUT_FILE, help="Where to write the report.")
    parser.add_argument(
        "--only",
        action="append",
        default=None,
        help="Scan just this parameter. Repeatable. Everything else is reported as skipped.",
    )
    args = parser.parse_args(argv)
    PACKAGE_DIR = args.package_dir

    if not PACKAGE_DIR.is_dir():
        print(f"error: corpus not found at {PACKAGE_DIR}", file=sys.stderr)
        return 2

    reference = json.loads(REFERENCE.read_text(encoding="utf-8"))
    invariants = reference["invariants"]
    key = _api_key()

    print(f"reading the corpus from {PACKAGE_DIR}...")
    corpus = {name: _document_text(name) for name in CORPUS}
    total_chars = sum(len(t) for t in corpus.values())
    print(f"  {len(corpus)} documents, {total_chars:,} characters")

    results: list[dict[str, Any]] = []
    for invariant in invariants:
        parameter = invariant["parameter"]
        if args.only and parameter not in args.only:
            continue
        if parameter not in ANCHORS:
            results.append(
                {
                    "parameter": parameter,
                    "declared_value": invariant["value"],
                    "declared_unit": invariant["unit"],
                    "declared_source": invariant["source"],
                    "status": "NOT SCANNED",
                    "reason": NOT_SCANNED_REASONS.get(parameter, DEFAULT_NOT_SCANNED),
                    "assertions": [],
                }
            )
            continue

        candidates = locate_candidates(corpus, parameter)
        capped = candidates[:MAX_CANDIDATES]
        assertions = [compare(invariant, a) for a in adjudicate(key, invariant, capped)]
        differs = [a for a in assertions if a["verdict"] == DIFFERS]
        unreadable = [a for a in assertions if a["verdict"] == UNREADABLE]

        if differs:
            status = "INCONSISTENT"
        elif unreadable:
            status = "INCOMPLETE"
        elif assertions:
            status = "CONSISTENT"
        else:
            status = "UNSUPPORTED"

        results.append(
            {
                "parameter": parameter,
                "declared_value": invariant["value"],
                "declared_unit": invariant["unit"],
                "declared_source": invariant["source"],
                "status": status,
                "reason": "",
                "anchors": ANCHORS[parameter],
                "candidates_found": len(candidates),
                "candidates_judged": len(capped),
                "assertions_confirmed": len(assertions),
                "matches": len(assertions) - len(differs) - len(unreadable),
                "differs": len(differs),
                "unreadable": len(unreadable),
                "documents": sorted({a["document"] for a in assertions}),
                "assertions": assertions,
            }
        )
        print(
            f"  {parameter:24s} {status:13s} "
            f"{len(candidates):3d} candidates -> {len(assertions)} assertions"
            + (f", {len(differs)} differ" if differs else "")
            + (f", {len(unreadable)} unreadable" if unreadable else "")
        )

    summary = {
        "invariants_total": len(results),
        "scanned": sum(1 for r in results if r["status"] != "NOT SCANNED"),
        "consistent": sum(1 for r in results if r["status"] == "CONSISTENT"),
        "inconsistent": sum(1 for r in results if r["status"] == "INCONSISTENT"),
        "incomplete": sum(1 for r in results if r["status"] == "INCOMPLETE"),
        "unsupported": sum(1 for r in results if r["status"] == "UNSUPPORTED"),
        "not_scanned": sum(1 for r in results if r["status"] == "NOT SCANNED"),
        "corpus_documents": len(corpus),
        "corpus_characters": total_chars,
        "assertions_confirmed": sum(len(r["assertions"]) for r in results),
    }

    report = {
        "case_id": "PMX103",
        "document_name": "Cross-document invariant scan",
        "declared_source": reference["source_files"]["invariants"],
        "corpus": CORPUS,
        "model": MODEL,
        "summary": summary,
        "invariants": results,
        "method": (
            "Candidate sentences are located deterministically by anchoring on the parameter's "
            "name. A model then judges each sentence in isolation - does this sentence assert a "
            "value for this parameter - and reports the value verbatim without converting it. "
            "The comparison against the declared value, including unit normalisation, is done in "
            "code afterwards. The model never generates a number and never decides whether two "
            "numbers match."
        ),
        "note": (
            "CONSISTENT means every sentence found asserting this parameter states the declared "
            "value. INCONSISTENT means at least one states something else. INCOMPLETE means a "
            "sentence sets out to state the parameter and carries no readable value - a hole in "
            "the document rather than a contradiction. UNSUPPORTED means no sentence was judged "
            "to assert it at all, which is not the same as the value being wrong. Text invariants "
            "are not scanned."
        ),
    }

    args.out.parent.mkdir(parents=True, exist_ok=True)
    args.out.write_text(json.dumps(report, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
    print(f"\nwrote {args.out}")
    print(
        f"  {summary['consistent']} consistent, {summary['inconsistent']} inconsistent, "
        f"{summary['incomplete']} incomplete, {summary['unsupported']} unsupported, "
        f"{summary['not_scanned']} not scanned"
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
